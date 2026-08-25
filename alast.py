
import io
import os
import re
import csv
import json
import shutil
import sqlite3
from pathlib import Path
from datetime import datetime, date

import streamlit as st
import pandas as pd

# =========================================================
# MIDO - CONSOLIDATED BUSINESS SYSTEM (NO AI YET)
# =========================================================

APP_NAME = "MIDO"
DB_FILE = Path("mido_v2.db")
ACCESS_DB_FILE = Path("mido_access_structure_only.db")
UPLOAD_ROOT = Path("uploads")

FACTORY_UPLOAD_ROOT = UPLOAD_ROOT / "factories"
ORDER_UPLOAD_ROOT = UPLOAD_ROOT / "orders"
SHIPMENT_UPLOAD_ROOT = UPLOAD_ROOT / "shipments"
SHIPMENT_INBOX_ROOT = UPLOAD_ROOT / "shipment_inbox"
PAYMENT_UPLOAD_ROOT = UPLOAD_ROOT / "payments"

st.set_page_config(page_title=APP_NAME, page_icon="📦", layout="wide")

# =========================================================
# BASIC STYLE
# =========================================================
st.markdown("""
<style>
.block-container {padding-top: 1.2rem; padding-bottom: 2rem;}
div[data-testid="stMetric"] {
    background: #f7f7f8;
    border: 1px solid #e5e7eb;
    padding: 12px 16px;
    border-radius: 14px;
}
.mido-card {
    border: 1px solid #e5e7eb;
    border-radius: 14px;
    padding: 14px;
    background: white;
    margin-bottom: 10px;
}
.small-muted {color:#6b7280;font-size:0.9rem;}
</style>
""", unsafe_allow_html=True)

# =========================================================
# DATABASE HELPERS
# =========================================================
def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def table_exists(table_name):
    conn = get_conn()
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table_name,)
    ).fetchone()
    conn.close()
    return row is not None

def get_columns(table_name):
    conn = get_conn()
    rows = conn.execute(f'PRAGMA table_info("{table_name}")').fetchall()
    conn.close()
    return {r["name"] for r in rows}

def add_missing_columns(table_name, definitions):
    if not table_exists(table_name):
        return
    existing = get_columns(table_name)
    conn = get_conn()
    for definition in definitions:
        col = definition.split()[0]
        if col not in existing:
            conn.execute(f'ALTER TABLE "{table_name}" ADD COLUMN {definition}')
    conn.commit()
    conn.close()

def execute(query, params=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(query, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id

def fetchall(query, params=()):
    conn = get_conn()
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return rows

def fetchone(query, params=()):
    conn = get_conn()
    row = conn.execute(query, params).fetchone()
    conn.close()
    return row

def init_database():
    conn = get_conn()
    cur = conn.cursor()

    cur.executescript("""
    CREATE TABLE IF NOT EXISTS app_settings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        setting_key TEXT UNIQUE NOT NULL,
        setting_value TEXT
    );

    CREATE TABLE IF NOT EXISTS factories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_name TEXT NOT NULL,
        chinese_name TEXT,
        brand_name TEXT,
        contact_person TEXT,
        phone TEXT,
        wechat TEXT,
        email TEXT,
        address TEXT,
        website TEXT,
        bank_info TEXT,
        payment_terms TEXT,
        monthly_order_deadline INTEGER DEFAULT 20,
        private_notes TEXT,
        public_notes TEXT,
        active INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS factory_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_id INTEGER NOT NULL,
        document_type TEXT,
        file_name TEXT NOT NULL,
        file_path TEXT,
        dropbox_path TEXT,
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS orders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_id INTEGER NOT NULL,
        pi_no TEXT,
        order_date TEXT,
        production_due_date TEXT,
        order_status TEXT DEFAULT 'Draft',
        currency TEXT DEFAULT 'USD',
        pi_amount REAL DEFAULT 0,
        expected_containers REAL DEFAULT 0,
        ordered_containers REAL DEFAULT 0,
        remaining_containers REAL DEFAULT 0,
        shipping_estimate REAL DEFAULT 0,
        total_estimated_amount REAL DEFAULT 0,
        destination TEXT,
        private_notes TEXT,
        public_notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS order_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        order_id INTEGER NOT NULL,
        size TEXT,
        pattern TEXT,
        description TEXT,
        quantity REAL DEFAULT 0,
        unit_price REAL DEFAULT 0,
        amount REAL DEFAULT 0,
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS order_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        order_id INTEGER NOT NULL,
        document_type TEXT,
        file_name TEXT NOT NULL,
        file_path TEXT,
        dropbox_path TEXT,
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shipments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_id INTEGER NOT NULL,
        order_id INTEGER,
        shipment_no TEXT,
        bl_no TEXT,
        shipping_company TEXT,
        destination TEXT,
        container_type TEXT DEFAULT '40HQ',
        containers_count REAL DEFAULT 0,
        ci_amount REAL DEFAULT 0,
        shipping_cost REAL DEFAULT 0,
        total_amount REAL DEFAULT 0,
        etd TEXT,
        eta TEXT,
        shipment_status TEXT DEFAULT 'Preparing',
        payment_status TEXT DEFAULT 'Not Paid',
        private_notes TEXT,
        public_notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shipment_containers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        shipment_id INTEGER NOT NULL,
        container_no TEXT,
        seal_no TEXT,
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shipment_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        shipment_id INTEGER NOT NULL,
        document_type TEXT,
        file_name TEXT NOT NULL,
        file_path TEXT,
        dropbox_path TEXT,
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shipment_intakes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_id INTEGER NOT NULL,
        order_id INTEGER,
        intake_name TEXT,
        status TEXT DEFAULT 'Uploaded',
        private_notes TEXT,
        extracted_pi_no TEXT,
        extracted_bl_no TEXT,
        extracted_shipping_company TEXT,
        extracted_destination TEXT,
        extracted_container_type TEXT,
        extracted_containers_count REAL DEFAULT 0,
        extracted_ci_amount REAL DEFAULT 0,
        extracted_shipping_cost REAL DEFAULT 0,
        extracted_etd TEXT,
        extracted_eta TEXT,
        extracted_container_numbers TEXT,
        analysis_notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shipment_intake_files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        intake_id INTEGER NOT NULL,
        file_name TEXT NOT NULL,
        file_path TEXT,
        dropbox_path TEXT,
        file_ext TEXT,
        detected_type TEXT DEFAULT 'Unclassified',
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS payments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        factory_id INTEGER,
        order_id INTEGER,
        shipment_id INTEGER,
        payment_date TEXT,
        amount REAL DEFAULT 0,
        currency TEXT DEFAULT 'USD',
        payment_type TEXT,
        bank_name TEXT,
        reference_no TEXT,
        supplier_confirmed INTEGER DEFAULT 0,
        private_notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS payment_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        payment_id INTEGER NOT NULL,
        document_type TEXT,
        file_name TEXT,
        file_path TEXT,
        dropbox_path TEXT,
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        phone TEXT,
        address TEXT,
        contact_person TEXT,
        credit_limit REAL DEFAULT 0,
        opening_balance REAL DEFAULT 0,
        private_notes TEXT,
        active INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS suppliers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        phone TEXT,
        address TEXT,
        contact_person TEXT,
        opening_balance REAL DEFAULT 0,
        private_notes TEXT,
        active INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS products (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sku TEXT,
        name TEXT NOT NULL,
        brand TEXT,
        category TEXT,
        unit TEXT,
        cost REAL DEFAULT 0,
        sale_price REAL DEFAULT 0,
        opening_qty REAL DEFAULT 0,
        min_qty REAL DEFAULT 0,
        notes TEXT,
        active INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS inventory_movements (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        movement_date TEXT,
        product_id INTEGER NOT NULL,
        movement_type TEXT,
        qty REAL DEFAULT 0,
        unit_cost REAL DEFAULT 0,
        reference_type TEXT,
        reference_id INTEGER,
        warehouse TEXT,
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS sales (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        invoice_no TEXT,
        sale_date TEXT,
        customer_id INTEGER,
        total_amount REAL DEFAULT 0,
        paid_amount REAL DEFAULT 0,
        currency TEXT DEFAULT 'USD',
        status TEXT DEFAULT 'Open',
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS sale_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sale_id INTEGER NOT NULL,
        product_id INTEGER,
        description TEXT,
        qty REAL DEFAULT 0,
        unit_price REAL DEFAULT 0,
        amount REAL DEFAULT 0
    );

    CREATE TABLE IF NOT EXISTS purchases (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        purchase_no TEXT,
        purchase_date TEXT,
        supplier_id INTEGER,
        factory_id INTEGER,
        total_amount REAL DEFAULT 0,
        paid_amount REAL DEFAULT 0,
        currency TEXT DEFAULT 'USD',
        status TEXT DEFAULT 'Open',
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS purchase_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        purchase_id INTEGER NOT NULL,
        product_id INTEGER,
        description TEXT,
        qty REAL DEFAULT 0,
        unit_price REAL DEFAULT 0,
        amount REAL DEFAULT 0
    );

    CREATE TABLE IF NOT EXISTS cash_transactions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        trans_date TEXT,
        trans_type TEXT,
        amount REAL DEFAULT 0,
        currency TEXT DEFAULT 'USD',
        party TEXT,
        reference TEXT,
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );
    """)

    conn.commit()
    conn.close()

    # Safe migrations from older MIDO versions
    for table, defs in {
        "factory_documents": ["dropbox_path TEXT"],
        "order_documents": ["dropbox_path TEXT"],
        "shipment_documents": ["dropbox_path TEXT"],
        "shipment_intake_files": ["dropbox_path TEXT"],
    }.items():
        add_missing_columns(table, defs)

init_database()

# =========================================================
# DROPBOX (OPTIONAL)
# =========================================================
def dropbox_token():
    try:
        return st.secrets.get("DROPBOX_ACCESS_TOKEN", "")
    except Exception:
        return ""

def dropbox_enabled():
    return bool(dropbox_token())

def upload_to_dropbox(local_path, remote_path):
    token = dropbox_token()
    if not token:
        return ""
    try:
        import dropbox
        dbx = dropbox.Dropbox(token)
        data = Path(local_path).read_bytes()
        dbx.files_upload(data, remote_path, mode=dropbox.files.WriteMode.overwrite)
        return remote_path
    except Exception:
        return ""

def storage_save(uploaded, local_dir, remote_dir):
    local_dir.mkdir(parents=True, exist_ok=True)
    safe_name = Path(uploaded.name).name
    target = local_dir / safe_name
    if target.exists():
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        target = local_dir / f"{target.stem}_{stamp}{target.suffix}"
    target.write_bytes(uploaded.getbuffer())
    remote_path = ""
    if dropbox_enabled():
        remote_path = upload_to_dropbox(target, f"{remote_dir}/{target.name}")
    return str(target), remote_path, target.name

# =========================================================
# FILE ANALYSIS (NO AI)
# =========================================================
def normalize_text(v):
    return re.sub(r"\s+", " ", str(v or "")).strip()

def read_file_text(path):
    path = Path(path)
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if ext in (".xlsx", ".xlsm"):
            from openpyxl import load_workbook
            wb = load_workbook(str(path), data_only=True, read_only=True)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(f"SHEET {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(v) for v in row if v is not None]
                    if vals:
                        chunks.append(" | ".join(vals))
            return "\n".join(chunks)
        if ext == ".csv":
            chunks = []
            with open(path, "r", encoding="utf-8-sig", errors="ignore", newline="") as f:
                for row in csv.reader(f):
                    chunks.append(" | ".join(row))
            return "\n".join(chunks)
        if ext == ".docx":
            from docx import Document
            doc = Document(str(path))
            chunks = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    chunks.append(" | ".join(c.text for c in row.cells))
            return "\n".join(chunks)
        if ext == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        pass
    return ""

def detect_document_type(file_name, text):
    s = (file_name + " " + text[:5000]).lower()
    if "packing list" in s or "packing" in s:
        return "Packing List"
    if "commercial invoice" in s or ("invoice" in s and "proforma" not in s):
        return "Commercial Invoice"
    if "certificate of origin" in s or "country of origin" in s:
        return "Certificate of Origin (CO)"
    if re.search(r"\bcoc\b", s) or "certificate of conformity" in s:
        return "COC"
    if "bill of lading" in s or re.search(r"\bb/?l\b", s):
        return "Bill of Lading"
    if "proforma" in s or re.search(r"\bpi\b", s):
        return "PI / Proforma Invoice"
    return "Unclassified"

def find_labeled_value(text, labels, token_pattern=r"[^|\n]{1,80}"):
    for label in labels:
        m = re.search(rf"{label}\s*[:#.\-]*\s*({token_pattern})", text, flags=re.I)
        if m:
            return normalize_text(m.group(1)).strip(" :|-_")
    return ""

def find_amount_after_labels(text, labels):
    for label in labels:
        m = re.search(rf"{label}[^\d$]{{0,30}}\$?\s*([\d,]+(?:\.\d+)?)", text, flags=re.I)
        if m:
            try:
                return float(m.group(1).replace(",", ""))
            except Exception:
                pass
    return 0.0

def find_date_after_label(text, label):
    m = re.search(
        rf"{label}[^0-9]{{0,25}}((?:20\d{{2}}[./-]\d{{1,2}}[./-]\d{{1,2}})|(?:\d{{1,2}}[./-]\d{{1,2}}[./-]20\d{{2}}))",
        text, flags=re.I
    )
    return m.group(1).replace(".", "-").replace("/", "-") if m else ""

def analyze_intake_files(intake_id):
    rows = fetchall("SELECT * FROM shipment_intake_files WHERE intake_id=? ORDER BY id", (intake_id,))
    all_text, all_names, notes = [], [], []

    for f in rows:
        path = Path(f["file_path"] or "")
        text = read_file_text(path)
        all_text.append(text)
        all_names.append(f["file_name"] or "")
        detected = detect_document_type(f["file_name"] or "", text)
        execute("UPDATE shipment_intake_files SET detected_type=? WHERE id=?", (detected, f["id"]))
        if not text and path.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
            notes.append(f"Image {f['file_name']} saved. Image-reading will be enabled when AI is added.")
        elif not text and path.suffix.lower() in (".xls", ".doc"):
            notes.append(f"Old-format file {f['file_name']} saved. XLSX/DOCX is preferred for automatic reading.")

    searchable = "\n".join(all_text + all_names)
    upper = searchable.upper()

    pi_no = find_labeled_value(searchable, [r"PI\s*(?:NO|NUMBER)", r"PROFORMA\s+INVOICE\s*(?:NO|NUMBER)"], r"[A-Z0-9][A-Z0-9._/-]{2,40}")
    bl_no = find_labeled_value(searchable, [r"B/?L\s*(?:NO|NUMBER)?", r"BILL\s+OF\s+LADING\s*(?:NO|NUMBER)?"], r"[A-Z0-9][A-Z0-9._/-]{5,40}")

    shipping_company = ""
    for company in ["MSC", "HPL", "CMA", "OOCL", "WANHAI", "ASYAD", "PIL", "MAERSK", "COSCO", "EVERGREEN"]:
        if re.search(rf"\b{re.escape(company)}\b", upper):
            shipping_company = company
            break

    destination = ""
    for canonical, variants in [
        ("UMM QASR", ["UMM QASR", "UMM,IRAQ", "UMM IRAQ"]),
        ("AQABA", ["AQABA", "AQABA,JOR", "AQABA JOR"]),
        ("MERSIN", ["MERSIN", "MERSIN,TUR", "MERSIN TUR"]),
    ]:
        if any(v in upper for v in variants):
            destination = canonical
            break

    container_numbers = sorted(set(re.findall(r"\b[A-Z]{4}\d{7}\b", upper)))
    containers_count = 0.0
    for pat in [
        r"(\d+(?:\.\d+)?)\s*[*X×]\s*40HQ",
        r"CONTS?\s*[*X×]?\s*40HQ[^\d]{0,20}(\d+(?:\.\d+)?)",
        r"(\d+(?:\.\d+)?)\s*(?:CONTAINER|CONTAINERS|CNTR|CNTRS)\b",
    ]:
        m = re.search(pat, upper, flags=re.I)
        if m:
            try:
                containers_count = float(m.group(1))
                break
            except Exception:
                pass
    if not containers_count and container_numbers:
        containers_count = float(len(container_numbers))

    ctype = "40HQ" if "40HQ" in upper else ("40GP" if "40GP" in upper else ("20GP" if "20GP" in upper else "40HQ"))
    ci_amount = find_amount_after_labels(searchable, [r"CI\s*AMOUNT(?:/USD)?", r"COMMERCIAL\s+INVOICE\s+AMOUNT", r"TOTAL\s+INVOICE\s+AMOUNT"])
    shipping_cost = find_amount_after_labels(searchable, [r"SHIPPING(?:\s+COST)?", r"FREIGHT"])
    etd = find_date_after_label(searchable, r"ETD")
    eta = find_date_after_label(searchable, r"ETA")

    execute("""
        UPDATE shipment_intakes
        SET extracted_pi_no=?, extracted_bl_no=?, extracted_shipping_company=?,
            extracted_destination=?, extracted_container_type=?, extracted_containers_count=?,
            extracted_ci_amount=?, extracted_shipping_cost=?, extracted_etd=?, extracted_eta=?,
            extracted_container_numbers=?, analysis_notes=?, status='Analyzed'
        WHERE id=?
    """, (
        pi_no, bl_no, shipping_company, destination, ctype, containers_count,
        ci_amount, shipping_cost, etd, eta, "\n".join(container_numbers),
        "\n".join(notes), intake_id
    ))

# =========================================================
# PDF REPORT
# =========================================================
def shipments_pdf(rows, title="MIDO - Main Shipment Report"):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics

    font = "Helvetica"
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("DejaVu", font_path))
            font = "DejaVu"
        except Exception:
            pass

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        rightMargin=18, leftMargin=18, topMargin=22, bottomMargin=18
    )
    styles = getSampleStyleSheet()
    styles["Title"].fontName = font
    story = [Paragraph(title, styles["Title"]), Spacer(1, 10)]

    header = [
        "PI No.", "Factory", "BL", "Ship Co.", "Destination",
        "Conts", "Type", "CI Amount", "Shipping", "ETD", "ETA", "State", "Payment"
    ]
    data = [header]
    for r in rows:
        data.append([
            str(r.get("PI No.", "")),
            str(r.get("Factory", "")),
            str(r.get("BL", "")),
            str(r.get("Ship Com.", "")),
            str(r.get("Destination", "")),
            str(r.get("CONTS", "")),
            str(r.get("Type", "")),
            f"{float(r.get('CI Amount/USD', 0) or 0):,.2f}",
            f"{float(r.get('Shipping', 0) or 0):,.2f}",
            str(r.get("ETD", "")),
            str(r.get("ETA", "")),
            str(r.get("State", "")),
            str(r.get("Payment", "")),
        ])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), font),
        ("FONTSIZE", (0,0), (-1,-1), 7),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E5E7EB")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.black),
        ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#6B7280")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN", (5,1), (8,-1), "RIGHT"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#FFFDE7")]),
    ]))
    story.append(table)
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()

# =========================================================
# LEGACY ACCESS DB
# =========================================================
def access_db_available():
    return ACCESS_DB_FILE.exists()

def get_access_conn():
    conn = sqlite3.connect(ACCESS_DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def access_table_names():
    if not access_db_available():
        return []
    conn = get_access_conn()
    rows = conn.execute("""
        SELECT name FROM sqlite_master
        WHERE type='table' AND name NOT LIKE 'sqlite_%'
        ORDER BY name
    """).fetchall()
    conn.close()
    return [r["name"] for r in rows if not r["name"].startswith("_")]

def access_table_columns(t):
    conn = get_access_conn()
    rows = conn.execute(f'PRAGMA table_info("{t}")').fetchall()
    conn.close()
    return [r["name"] for r in rows]

def access_table_count(t):
    conn = get_access_conn()
    n = conn.execute(f'SELECT COUNT(*) AS n FROM "{t}"').fetchone()["n"]
    conn.close()
    return n

# =========================================================
# COMMON HELPERS
# =========================================================
def factory_map(active_only=True):
    q = "SELECT id, factory_name, brand_name FROM factories"
    if active_only:
        q += " WHERE COALESCE(active,1)=1"
    q += " ORDER BY factory_name"
    rows = fetchall(q)
    return {
        f"{r['factory_name']} — {r['brand_name'] or 'بدون براند'}": r["id"]
        for r in rows
    }

def safe_float(v):
    try:
        return float(v or 0)
    except Exception:
        return 0.0

# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.title("📦 MIDO")
st.sidebar.caption("Business & Shipment Management System")
st.sidebar.write("Dropbox:", "✅" if dropbox_enabled() else "Local storage")

page = st.sidebar.radio(
    "القائمة الرئيسية",
    [
        "الرئيسية",
        "المعامل",
        "الطلبيات",
        "المشحونات",
        "الدفعات",
        "العملاء",
        "المواد والمخزون",
        "المبيعات",
        "المشتريات",
        "الصندوق",
        "المستندات",
        "نظام Access القديم",
        "التقارير",
    ],
)

# =========================================================
# HOME
# =========================================================
if page == "الرئيسية":
    st.title("📦 MIDO")
    st.subheader("نظام إدارة الشركة والطلبيات والمشحونات")

    factory_count = fetchone("SELECT COUNT(*) AS n FROM factories")["n"]
    order_count = fetchone("SELECT COUNT(*) AS n FROM orders")["n"]
    shipment_count = fetchone("SELECT COUNT(*) AS n FROM shipments")["n"]
    in_transit = fetchone("""
        SELECT COALESCE(SUM(containers_count),0) AS n
        FROM shipments
        WHERE shipment_status IN ('Booked','Shipped','In Transit')
    """)["n"]
    unpaid = fetchone("""
        SELECT COUNT(*) AS n FROM shipments
        WHERE payment_status NOT IN ('Paid','Supplier Confirmed')
    """)["n"]

    c1,c2,c3,c4,c5 = st.columns(5)
    c1.metric("المعامل", factory_count)
    c2.metric("الطلبيات", order_count)
    c3.metric("المشحونات", shipment_count)
    c4.metric("الحاويات بالطريق", f"{safe_float(in_transit):g}")
    c5.metric("شحنات غير مكتملة الدفع", unpaid)

    st.subheader("الشحنات القادمة")
    rows = fetchall("""
        SELECT s.*, f.factory_name, o.pi_no
        FROM shipments s
        LEFT JOIN factories f ON f.id=s.factory_id
        LEFT JOIN orders o ON o.id=s.order_id
        WHERE s.shipment_status NOT IN ('Delivered','Cancelled')
        ORDER BY CASE WHEN s.eta IS NULL OR s.eta='' THEN 1 ELSE 0 END, s.eta
        LIMIT 20
    """)
    if rows:
        st.dataframe([{
            "Factory":r["factory_name"] or "-",
            "PI":r["pi_no"] or "-",
            "BL":r["bl_no"] or "-",
            "Containers":r["containers_count"] or 0,
            "ETD":r["etd"] or "-",
            "ETA":r["eta"] or "-",
            "Status":r["shipment_status"] or "-",
            "Payment":r["payment_status"] or "-",
        } for r in rows], use_container_width=True, hide_index=True)
    else:
        st.info("لا توجد شحنات حالياً.")

    st.subheader("متابعة طلبات المصانع قبل يوم 20")
    facs = fetchall("SELECT factory_name, monthly_order_deadline FROM factories WHERE COALESCE(active,1)=1 ORDER BY factory_name")
    if facs:
        st.dataframe([{
            "المعمل":r["factory_name"],
            "موعد تجهيز الطلبية":f"قبل يوم {r['monthly_order_deadline'] or 20} من كل شهر"
        } for r in facs], use_container_width=True, hide_index=True)

# =========================================================
# FACTORIES
# =========================================================
elif page == "المعامل":
    st.title("🏭 المعامل")
    t1,t2 = st.tabs(["➕ إضافة معمل","📋 المعامل المسجلة"])

    with t1:
        with st.form("add_factory", clear_on_submit=True):
            a,b,c = st.columns(3)
            name = a.text_input("اسم المعمل *")
            chinese = b.text_input("الاسم بالصيني")
            brand = c.text_input("البراند")
            contact = a.text_input("الشخص المسؤول")
            phone = b.text_input("الهاتف")
            wechat = c.text_input("WeChat")
            email = a.text_input("Email")
            website = b.text_input("Website")
            deadline = c.number_input("موعد الطلب الشهري",1,31,20)
            address = st.text_area("العنوان")
            bank = st.text_area("معلومات الحساب البنكي")
            terms = st.text_area("شروط الدفع")
            n1,n2 = st.columns(2)
            private = n1.text_area("🔒 ملاحظات خاصة")
            public = n2.text_area("📝 ملاحظات عامة")
            if st.form_submit_button("💾 حفظ المعمل", type="primary"):
                if not name.strip():
                    st.error("اسم المعمل مطلوب.")
                else:
                    execute("""
                        INSERT INTO factories
                        (factory_name,chinese_name,brand_name,contact_person,phone,wechat,email,address,website,
                         bank_info,payment_terms,monthly_order_deadline,private_notes,public_notes,active)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,1)
                    """,(name.strip(),chinese.strip(),brand.strip(),contact.strip(),phone.strip(),wechat.strip(),
                         email.strip(),address.strip(),website.strip(),bank.strip(),terms.strip(),int(deadline),
                         private.strip(),public.strip()))
                    st.success("تم حفظ المعمل.")

    with t2:
        factories = fetchall("SELECT * FROM factories ORDER BY active DESC, factory_name")
        search = st.text_input("🔎 بحث عن معمل", key="factory_search")
        for f in factories:
            hay = " ".join([str(f["factory_name"] or ""),str(f["brand_name"] or ""),str(f["contact_person"] or "")]).lower()
            if search and search.lower() not in hay:
                continue
            title = f"{'🟢' if f['active'] else '⚪'} {f['factory_name']} — {f['brand_name'] or '-'}"
            with st.expander(title):
                c1,c2,c3 = st.columns(3)
                c1.write(f"**Contact:** {f['contact_person'] or '-'}")
                c1.write(f"**Phone:** {f['phone'] or '-'}")
                c1.write(f"**WeChat:** {f['wechat'] or '-'}")
                c2.write(f"**Email:** {f['email'] or '-'}")
                c2.write(f"**Website:** {f['website'] or '-'}")
                c2.write(f"**Monthly deadline:** {f['monthly_order_deadline'] or 20}")
                c3.write(f"**Bank:** {f['bank_info'] or '-'}")
                c3.write(f"**Payment terms:** {f['payment_terms'] or '-'}")

                n1,n2 = st.columns(2)
                n1.info(f["private_notes"] or "لا توجد ملاحظات خاصة")
                n2.info(f["public_notes"] or "لا توجد ملاحظات عامة")

                st.markdown("#### 📁 ملفات المعمل")
                dtype = st.selectbox("نوع الملف",[
                    "Company Registration","Contract / Agreement","Agency / Authorization",
                    "Bank Details","Price List","Certificate","PI / Quotation","Other"
                ], key=f"fdtype_{f['id']}")
                files = st.file_uploader("ارفع ملف أو أكثر", accept_multiple_files=True,
                    type=["pdf","xlsx","xls","docx","doc","png","jpg","jpeg","webp"],
                    key=f"fup_{f['id']}")
                if files and st.button("حفظ الملفات", key=f"fsave_{f['id']}"):
                    for up in files:
                        lp,dp,fn = storage_save(
                            up,
                            FACTORY_UPLOAD_ROOT / str(f["id"]),
                            f"/MIDO/Factories/{f['factory_name']}/Factory Documents"
                        )
                        execute("""INSERT INTO factory_documents
                            (factory_id,document_type,file_name,file_path,dropbox_path)
                            VALUES (?,?,?,?,?)""",(f["id"],dtype,fn,lp,dp))
                    st.success("تم حفظ الملفات.")

# =========================================================
# ORDERS
# =========================================================
elif page == "الطلبيات":
    st.title("📋 الطلبيات")
    fmap = factory_map()
    if not fmap:
        st.warning("أضف معمل أولاً.")
    else:
        t1,t2 = st.tabs(["➕ طلبية جديدة","📋 الطلبيات المسجلة"])
        with t1:
            with st.form("new_order", clear_on_submit=True):
                a,b,c = st.columns(3)
                flabel = a.selectbox("المعمل", list(fmap.keys()))
                pi = b.text_input("PI No. *")
                status = c.selectbox("الحالة",["Draft","Confirmed","In Production","Ready","Partially Shipped","Completed","Cancelled"])
                a,b,c = st.columns(3)
                od = a.date_input("تاريخ الطلبية", date.today())
                due = b.date_input("موعد الإنتاج", date.today())
                curr = c.selectbox("العملة",["USD","CNY","EUR"])
                a,b,c = st.columns(3)
                pia = a.number_input("PI Amount",0.0,step=100.0)
                expc = b.number_input("الحاويات المطلوبة",0.0,step=0.5)
                confc = c.number_input("الحاويات المؤكدة",0.0,step=0.5)
                a,b,c = st.columns(3)
                ship_est = a.number_input("Shipping Estimate",0.0,step=100.0)
                dest = b.text_input("Destination")
                total = c.number_input("Total Estimate",0.0,step=100.0)
                n1,n2 = st.columns(2)
                private=n1.text_area("🔒 ملاحظات خاصة")
                public=n2.text_area("📝 ملاحظات عامة")
                if st.form_submit_button("💾 حفظ الطلبية",type="primary"):
                    if not pi.strip():
                        st.error("PI No. مطلوب.")
                    else:
                        rem=max(expc-confc,0)
                        execute("""INSERT INTO orders
                            (factory_id,pi_no,order_date,production_due_date,order_status,currency,pi_amount,
                             expected_containers,ordered_containers,remaining_containers,shipping_estimate,
                             total_estimated_amount,destination,private_notes,public_notes)
                             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                             (fmap[flabel],pi.strip(),str(od),str(due),status,curr,pia,expc,confc,rem,
                              ship_est,total,dest.strip(),private.strip(),public.strip()))
                        st.success("تم حفظ الطلبية.")

        with t2:
            orders=fetchall("""SELECT o.*,f.factory_name,f.brand_name FROM orders o
                LEFT JOIN factories f ON f.id=o.factory_id ORDER BY o.id DESC""")
            for o in orders:
                with st.expander(f"#{o['id']} — {o['pi_no'] or '-'} — {o['factory_name'] or '-'}"):
                    m1,m2,m3,m4=st.columns(4)
                    m1.metric("PI Amount",f"{safe_float(o['pi_amount']):,.2f} {o['currency'] or 'USD'}")
                    m2.metric("Required",f"{safe_float(o['expected_containers']):g}")
                    m3.metric("Confirmed",f"{safe_float(o['ordered_containers']):g}")
                    m4.metric("Left",f"{safe_float(o['remaining_containers']):g}")

                    st.markdown("#### تفاصيل الطلبية")
                    with st.form(f"oi_{o['id']}", clear_on_submit=True):
                        a,b,c=st.columns(3)
                        size=a.text_input("Size")
                        pattern=b.text_input("Pattern")
                        desc=c.text_input("Description")
                        a,b,c=st.columns(3)
                        qty=a.number_input("Qty",0.0,step=1.0)
                        price=b.number_input("Unit Price",0.0,step=0.01)
                        note=c.text_input("Note")
                        if st.form_submit_button("➕ إضافة سطر"):
                            execute("""INSERT INTO order_items
                                (order_id,size,pattern,description,quantity,unit_price,amount,notes)
                                VALUES (?,?,?,?,?,?,?,?)""",
                                (o["id"],size.strip(),pattern.strip(),desc.strip(),qty,price,qty*price,note.strip()))
                            st.rerun()

                    items=fetchall("SELECT * FROM order_items WHERE order_id=? ORDER BY id",(o["id"],))
                    if items:
                        st.dataframe([dict(x) for x in items],use_container_width=True,hide_index=True)

                    st.markdown("#### 📁 ملفات الطلبية")
                    dtype=st.selectbox("نوع الملف",["PI / Proforma Invoice","Order Excel","Order PDF","Price List","Specification","Confirmation","Other"],key=f"odt_{o['id']}")
                    files=st.file_uploader("ارفع ملفات الطلبية",accept_multiple_files=True,
                        type=["pdf","xlsx","xls","csv","docx","doc","png","jpg","jpeg","webp"],
                        key=f"oup_{o['id']}")
                    if files and st.button("حفظ ملفات الطلبية",key=f"osave_{o['id']}"):
                        for up in files:
                            lp,dp,fn=storage_save(up,ORDER_UPLOAD_ROOT/str(o["id"]),
                                f"/MIDO/Factories/{o['factory_name']}/Orders/{o['pi_no'] or o['id']}")
                            execute("""INSERT INTO order_documents
                                (order_id,document_type,file_name,file_path,dropbox_path)
                                VALUES (?,?,?,?,?)""",(o["id"],dtype,fn,lp,dp))
                        st.success("تم حفظ الملفات.")

# =========================================================
# SHIPMENTS
# =========================================================
elif page == "المشحونات":
    st.title("🚢 المشحونات")
    fmap=factory_map()
    if not fmap:
        st.warning("أضف معمل أولاً.")
    else:
        t1,t2,t3=st.tabs(["📥 إضافة من الملفات","✍️ إدخال يدوي","📋 المشحونات الرئيسية"])

        # FILE-FIRST
        with t1:
            st.subheader("إضافة شحنة من ملفاتها الأصلية")
            flabel=st.selectbox("المعمل",list(fmap.keys()),key="intake_factory")
            fid=fmap[flabel]
            orders=fetchall("SELECT id,pi_no FROM orders WHERE factory_id=? ORDER BY id DESC",(fid,))
            omap={"بدون ربط بطلبية":None}
            for r in orders:
                omap[f"{r['pi_no'] or 'Order'} — #{r['id']}"]=r["id"]
            olabel=st.selectbox("الطلبية المرتبطة",list(omap.keys()),key="intake_order")
            name=st.text_input("اسم مؤقت للشحنة",placeholder="مثال: BOTRIAN AUG-2026 / Shipment 01")
            files=st.file_uploader("ارفع جميع الملفات دفعة واحدة",accept_multiple_files=True,
                type=["pdf","xlsx","xls","csv","docx","doc","png","jpg","jpeg","webp"],
                key="intake_files")
            private=st.text_area("🔒 ملاحظات خاصة",key="intake_private")
            if files and st.button("📥 حفظ الشحنة وملفاتها",type="primary"):
                iid=execute("""INSERT INTO shipment_intakes
                    (factory_id,order_id,intake_name,status,private_notes)
                    VALUES (?,?,?,'Uploaded',?)""",
                    (fid,omap[olabel],name.strip() or f"Shipment {datetime.now():%Y-%m-%d %H:%M}",private.strip()))
                frow=fetchone("SELECT factory_name FROM factories WHERE id=?",(fid,))
                for up in files:
                    lp,dp,fn=storage_save(up,SHIPMENT_INBOX_ROOT/str(iid),
                        f"/MIDO/Factories/{frow['factory_name']}/Shipment Inbox/{iid}")
                    execute("""INSERT INTO shipment_intake_files
                        (intake_id,file_name,file_path,dropbox_path,file_ext,detected_type)
                        VALUES (?,?,?,?,?,'Unclassified')""",
                        (iid,fn,lp,dp,Path(fn).suffix.lower().replace(".","")))
                st.success("تم حفظ الشحنة وملفاتها.")
                st.rerun()

            intakes=fetchall("""SELECT i.*,f.factory_name,o.pi_no FROM shipment_intakes i
                LEFT JOIN factories f ON f.id=i.factory_id
                LEFT JOIN orders o ON o.id=i.order_id ORDER BY i.id DESC""")
            for i in intakes:
                with st.expander(f"#{i['id']} — {i['factory_name'] or '-'} — {i['intake_name'] or '-'}"):
                    fr=fetchall("SELECT * FROM shipment_intake_files WHERE intake_id=? ORDER BY id",(i["id"],))
                    if fr:
                        st.dataframe([{
                            "File":x["file_name"],"Type":(x["file_ext"] or "").upper(),
                            "Detected As":x["detected_type"] or "Unclassified"
                        } for x in fr],use_container_width=True,hide_index=True)
                    if st.button("⚙️ تحليل الملفات",key=f"an_{i['id']}",type="primary"):
                        analyze_intake_files(i["id"])
                        st.rerun()

                    analyzed=fetchone("SELECT * FROM shipment_intakes WHERE id=?",(i["id"],))
                    if analyzed and analyzed["status"] in ("Analyzed","Converted"):
                        st.markdown("#### البيانات المستخرجة — راجعها قبل الاعتماد")
                        with st.form(f"review_{i['id']}"):
                            a,b,c=st.columns(3)
                            epi=a.text_input("PI No.",value=analyzed["extracted_pi_no"] or "")
                            ebl=b.text_input("B/L No.",value=analyzed["extracted_bl_no"] or "")
                            eco=c.text_input("Shipping Company",value=analyzed["extracted_shipping_company"] or "")
                            a,b,c=st.columns(3)
                            edest=a.text_input("Destination",value=analyzed["extracted_destination"] or "")
                            ecnt=b.number_input("عدد الحاويات",0.0,step=0.5,value=safe_float(analyzed["extracted_containers_count"]))
                            ctypes=["40HQ","40GP","20GP","Other"]
                            current=analyzed["extracted_container_type"] or "40HQ"
                            etype=c.selectbox("نوع الحاوية",ctypes,index=ctypes.index(current) if current in ctypes else 0)
                            a,b=st.columns(2)
                            eci=a.number_input("CI Amount",0.0,step=100.0,value=safe_float(analyzed["extracted_ci_amount"]))
                            eship=b.number_input("Shipping Cost",0.0,step=100.0,value=safe_float(analyzed["extracted_shipping_cost"]))
                            a,b=st.columns(2)
                            eetd=a.text_input("ETD",value=analyzed["extracted_etd"] or "")
                            eeta=b.text_input("ETA",value=analyzed["extracted_eta"] or "")
                            econtainers=st.text_area("Container Numbers — رقم بكل سطر",value=analyzed["extracted_container_numbers"] or "")
                            status=st.selectbox("حالة الشحنة",["Preparing","Ready","Booked","Shipped","In Transit","Arrived","Customs","Delivered"])
                            if st.form_submit_button("✅ اعتماد وإنشاء الشحنة",type="primary"):
                                sid=execute("""INSERT INTO shipments
                                    (factory_id,order_id,shipment_no,bl_no,shipping_company,destination,container_type,
                                     containers_count,ci_amount,shipping_cost,total_amount,etd,eta,shipment_status,
                                     payment_status,private_notes,public_notes)
                                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,'Not Paid',?,'')""",
                                    (i["factory_id"],i["order_id"],i["intake_name"],ebl.strip(),eco.strip(),edest.strip(),etype,
                                     ecnt,eci,eship,eci+eship,eetd.strip(),eeta.strip(),status,i["private_notes"] or ""))

                                # Try matching order by PI
                                if not i["order_id"] and epi.strip():
                                    m=fetchone("SELECT id FROM orders WHERE factory_id=? AND pi_no=? ORDER BY id DESC LIMIT 1",(i["factory_id"],epi.strip()))
                                    if m:
                                        execute("UPDATE shipments SET order_id=? WHERE id=?",(m["id"],sid))

                                sdir=SHIPMENT_UPLOAD_ROOT/str(sid)
                                sdir.mkdir(parents=True,exist_ok=True)
                                for x in fr:
                                    source=Path(x["file_path"] or "")
                                    target=sdir/source.name
                                    if source.exists():
                                        shutil.copy2(source,target)
                                    execute("""INSERT INTO shipment_documents
                                        (shipment_id,document_type,file_name,file_path,dropbox_path)
                                        VALUES (?,?,?,?,?)""",
                                        (sid,x["detected_type"] or "Unclassified",x["file_name"],str(target),x["dropbox_path"] or ""))

                                for cn in [x.strip() for x in econtainers.splitlines() if x.strip()]:
                                    execute("""INSERT INTO shipment_containers
                                        (shipment_id,container_no,seal_no,notes) VALUES (?,?,'','')""",(sid,cn))

                                execute("""UPDATE shipment_intakes SET status='Converted',
                                    extracted_pi_no=?,extracted_bl_no=?,extracted_shipping_company=?,extracted_destination=?,
                                    extracted_container_type=?,extracted_containers_count=?,extracted_ci_amount=?,
                                    extracted_shipping_cost=?,extracted_etd=?,extracted_eta=?,extracted_container_numbers=?
                                    WHERE id=?""",(epi.strip(),ebl.strip(),eco.strip(),edest.strip(),etype,ecnt,eci,eship,eetd.strip(),eeta.strip(),econtainers.strip(),i["id"]))
                                st.success("تم إنشاء الشحنة وربط ملفاتها.")
                                st.rerun()
                        if analyzed["analysis_notes"]:
                            st.warning(analyzed["analysis_notes"])

        # MANUAL
        with t2:
            flabel=st.selectbox("المعمل",list(fmap.keys()),key="manual_factory")
            fid=fmap[flabel]
            ords=fetchall("SELECT id,pi_no FROM orders WHERE factory_id=? ORDER BY id DESC",(fid,))
            omap={"بدون طلبية":None}
            for r in ords:
                omap[f"{r['pi_no'] or 'Order'} — #{r['id']}"]=r["id"]
            with st.form("manual_ship",clear_on_submit=True):
                a,b,c=st.columns(3)
                olabel=a.selectbox("الطلبية",list(omap.keys()))
                sno=b.text_input("Shipment No.")
                bl=c.text_input("B/L No.")
                a,b,c=st.columns(3)
                shipco=a.text_input("Shipping Company")
                dest=b.text_input("Destination")
                ctype=c.selectbox("Container Type",["40HQ","40GP","20GP","Other"])
                a,b,c=st.columns(3)
                cnt=a.number_input("Containers",0.0,step=0.5)
                ci=b.number_input("CI Amount",0.0,step=100.0)
                sc=c.number_input("Shipping Cost",0.0,step=100.0)
                a,b,c=st.columns(3)
                etd=a.date_input("ETD",date.today())
                eta=b.date_input("ETA",date.today())
                status=c.selectbox("Status",["Preparing","Ready","Booked","Shipped","In Transit","Arrived","Customs","Delivered","Cancelled"])
                pay=st.selectbox("Payment",["Not Paid","Deposit Paid","Partially Paid","Paid","Supplier Confirmed"])
                n1,n2=st.columns(2)
                private=n1.text_area("🔒 Private Notes")
                public=n2.text_area("Public Notes")
                if st.form_submit_button("💾 حفظ الشحنة",type="primary"):
                    execute("""INSERT INTO shipments
                        (factory_id,order_id,shipment_no,bl_no,shipping_company,destination,container_type,containers_count,
                         ci_amount,shipping_cost,total_amount,etd,eta,shipment_status,payment_status,private_notes,public_notes)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (fid,omap[olabel],sno.strip(),bl.strip(),shipco.strip(),dest.strip(),ctype,cnt,ci,sc,ci+sc,
                         str(etd),str(eta),status,pay,private.strip(),public.strip()))
                    st.success("تم حفظ الشحنة.")

        # MAIN LIST
        with t3:
            shipments=fetchall("""SELECT s.*,f.factory_name,f.brand_name,o.pi_no
                FROM shipments s LEFT JOIN factories f ON f.id=s.factory_id
                LEFT JOIN orders o ON o.id=s.order_id ORDER BY s.id DESC""")
            if not shipments:
                st.info("لا توجد مشحونات.")
            else:
                a,b,c=st.columns(3)
                search=a.text_input("🔎 بحث",placeholder="PI / BL / Factory / Ship Co.")
                sf=b.selectbox("الحالة",["الكل","Preparing","Ready","Booked","Shipped","In Transit","Arrived","Customs","Delivered","Cancelled"])
                pf=c.selectbox("الدفع",["الكل","Not Paid","Deposit Paid","Partially Paid","Paid","Supplier Confirmed"])
                filtered=[]
                for s in shipments:
                    hay=" ".join([str(s["shipment_no"] or ""),str(s["bl_no"] or ""),str(s["pi_no"] or ""),
                                  str(s["factory_name"] or ""),str(s["shipping_company"] or "")]).lower()
                    if search and search.lower() not in hay: continue
                    if sf!="الكل" and s["shipment_status"]!=sf: continue
                    if pf!="الكل" and s["payment_status"]!=pf: continue
                    filtered.append(s)

                rows=[]
                for s in filtered:
                    rows.append({
                        "PI No.":s["pi_no"] or "-","Factory":s["factory_name"] or "-",
                        "BL":s["bl_no"] or "-","Ship Com.":s["shipping_company"] or "-",
                        "Destination":s["destination"] or "-","CONTS":s["containers_count"] or 0,
                        "Type":s["container_type"] or "40HQ","CI Amount/USD":s["ci_amount"] or 0,
                        "Shipping":s["shipping_cost"] or 0,"Total Amount":s["total_amount"] or 0,
                        "ETD":s["etd"] or "-","ETA":s["eta"] or "-",
                        "State":s["shipment_status"] or "-","Payment":s["payment_status"] or "-"
                    })

                m1,m2,m3,m4=st.columns(4)
                m1.metric("المشحونات",len(rows))
                m2.metric("الحاويات",f"{sum(safe_float(x['CONTS']) for x in rows):g}")
                m3.metric("CI Amount",f"${sum(safe_float(x['CI Amount/USD']) for x in rows):,.2f}")
                m4.metric("Shipping",f"${sum(safe_float(x['Shipping']) for x in rows):,.2f}")

                st.dataframe(rows,use_container_width=True,hide_index=True)
                pdf=shipments_pdf(rows)
                st.download_button("📄 تنزيل PDF للمشحونات الحالية",pdf,
                    file_name=f"MIDO_Shipments_{date.today()}.pdf",mime="application/pdf")

                st.divider()
                for s in filtered:
                    with st.expander(f"#{s['id']} — {s['factory_name'] or '-'} — {s['pi_no'] or '-'} — {s['bl_no'] or '-'}"):
                        a,b,c,d=st.columns(4)
                        a.metric("Containers",f"{safe_float(s['containers_count']):g}")
                        b.metric("CI",f"${safe_float(s['ci_amount']):,.2f}")
                        c.metric("Shipping",f"${safe_float(s['shipping_cost']):,.2f}")
                        d.metric("Total",f"${safe_float(s['total_amount']):,.2f}")
                        n1,n2=st.columns(2)
                        n1.info(s["private_notes"] or "No private notes")
                        n2.info(s["public_notes"] or "No public notes")

                        st.markdown("#### Container Numbers")
                        with st.form(f"cn_{s['id']}",clear_on_submit=True):
                            a,b,c=st.columns(3)
                            cn=a.text_input("Container No.")
                            seal=b.text_input("Seal No.")
                            note=c.text_input("Note")
                            if st.form_submit_button("Add Container"):
                                if cn.strip():
                                    execute("INSERT INTO shipment_containers (shipment_id,container_no,seal_no,notes) VALUES (?,?,?,?)",
                                            (s["id"],cn.strip(),seal.strip(),note.strip()))
                                    st.rerun()
                        conts=fetchall("SELECT * FROM shipment_containers WHERE shipment_id=? ORDER BY id",(s["id"],))
                        if conts:
                            st.dataframe([dict(x) for x in conts],use_container_width=True,hide_index=True)

                        st.markdown("#### Documents")
                        dtype=st.selectbox("Document Type",["Commercial Invoice","Packing List","Certificate of Origin (CO)","COC","Bill of Lading","QR Code","Other"],key=f"sdt_{s['id']}")
                        files=st.file_uploader("Upload files",accept_multiple_files=True,
                            type=["pdf","xlsx","xls","csv","docx","doc","png","jpg","jpeg","webp"],key=f"sup_{s['id']}")
                        if files and st.button("Save Documents",key=f"ssave_{s['id']}"):
                            for up in files:
                                lp,dp,fn=storage_save(up,SHIPMENT_UPLOAD_ROOT/str(s["id"]),
                                    f"/MIDO/Factories/{s['factory_name']}/Shipments/{s['shipment_no'] or s['id']}")
                                execute("""INSERT INTO shipment_documents
                                    (shipment_id,document_type,file_name,file_path,dropbox_path)
                                    VALUES (?,?,?,?,?)""",(s["id"],dtype,fn,lp,dp))
                            st.success("Documents saved.")

# =========================================================
# PAYMENTS
# =========================================================
elif page == "الدفعات":
    st.title("💰 الدفعات")
    fmap=factory_map(False)
    t1,t2=st.tabs(["➕ تسجيل دفعة","📋 سجل الدفعات"])
    with t1:
        if not fmap:
            st.warning("أضف معمل أولاً.")
        else:
            flabel=st.selectbox("المعمل",list(fmap.keys()),key="pay_factory")
            fid=fmap[flabel]
            ords=fetchall("SELECT id,pi_no FROM orders WHERE factory_id=? ORDER BY id DESC",(fid,))
            ships=fetchall("SELECT id,shipment_no,bl_no FROM shipments WHERE factory_id=? ORDER BY id DESC",(fid,))
            omap={"بدون طلبية":None}; smap={"بدون شحنة":None}
            for x in ords: omap[f"{x['pi_no'] or 'Order'} — #{x['id']}"]=x["id"]
            for x in ships: smap[f"{x['shipment_no'] or 'Shipment'} — {x['bl_no'] or 'No BL'}"]=x["id"]
            with st.form("pay_form",clear_on_submit=True):
                a,b,c=st.columns(3)
                olabel=a.selectbox("الطلبية",list(omap.keys()))
                slabel=b.selectbox("الشحنة",list(smap.keys()))
                pdate=c.date_input("تاريخ الدفع",date.today())
                a,b,c=st.columns(3)
                amount=a.number_input("المبلغ",0.0,step=100.0)
                curr=b.selectbox("العملة",["USD","CNY","EUR"])
                ptype=c.selectbox("نوع الدفع",["Deposit","Balance","Full Payment","Freight","Other"])
                a,b,c=st.columns(3)
                bank=a.text_input("Bank")
                ref=b.text_input("Reference No.")
                confirmed=c.checkbox("المعمل أكد وصول المبلغ")
                note=st.text_area("🔒 ملاحظات خاصة")
                if st.form_submit_button("💾 حفظ الدفعة",type="primary"):
                    pid=execute("""INSERT INTO payments
                        (factory_id,order_id,shipment_id,payment_date,amount,currency,payment_type,bank_name,reference_no,
                         supplier_confirmed,private_notes)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                        (fid,omap[olabel],smap[slabel],str(pdate),amount,curr,ptype,bank.strip(),ref.strip(),1 if confirmed else 0,note.strip()))
                    st.session_state["last_payment_id"]=pid
                    st.success("تم حفظ الدفعة.")

            pid=st.session_state.get("last_payment_id")
            if pid:
                files=st.file_uploader("ارفع Bank Slip",accept_multiple_files=True,
                    type=["pdf","png","jpg","jpeg","webp"],key="bank_slip")
                if files and st.button("حفظ Bank Slip"):
                    for up in files:
                        lp,dp,fn=storage_save(up,PAYMENT_UPLOAD_ROOT/str(pid),f"/MIDO/Payments/{pid}")
                        execute("""INSERT INTO payment_documents
                            (payment_id,document_type,file_name,file_path,dropbox_path)
                            VALUES (?,'Bank Slip',?,?,?)""",(pid,fn,lp,dp))
                    st.success("تم حفظ Bank Slip.")

    with t2:
        rows=fetchall("""SELECT p.*,f.factory_name,o.pi_no,s.bl_no
            FROM payments p LEFT JOIN factories f ON f.id=p.factory_id
            LEFT JOIN orders o ON o.id=p.order_id
            LEFT JOIN shipments s ON s.id=p.shipment_id
            ORDER BY p.id DESC""")
        if rows:
            st.dataframe([{
                "Date":r["payment_date"],"Factory":r["factory_name"],"PI":r["pi_no"],
                "BL":r["bl_no"],"Amount":r["amount"],"Currency":r["currency"],
                "Type":r["payment_type"],"Confirmed":"Yes" if r["supplier_confirmed"] else "No",
                "Reference":r["reference_no"]
            } for r in rows],use_container_width=True,hide_index=True)
        else:
            st.info("لا توجد دفعات.")

# =========================================================
# CUSTOMERS
# =========================================================
elif page == "العملاء":
    st.title("👥 العملاء")
    t1,t2=st.tabs(["➕ إضافة عميل","📋 العملاء"])
    with t1:
        with st.form("customer_add",clear_on_submit=True):
            a,b=st.columns(2)
            name=a.text_input("اسم العميل *")
            phone=b.text_input("الهاتف")
            a,b=st.columns(2)
            contact=a.text_input("الشخص المسؤول")
            limit=b.number_input("Credit Limit",0.0,step=100.0)
            address=st.text_area("العنوان")
            opening=st.number_input("Opening Balance",step=100.0)
            notes=st.text_area("🔒 ملاحظات خاصة")
            if st.form_submit_button("حفظ العميل",type="primary"):
                if name.strip():
                    execute("""INSERT INTO customers
                        (name,phone,address,contact_person,credit_limit,opening_balance,private_notes)
                        VALUES (?,?,?,?,?,?,?)""",(name.strip(),phone.strip(),address.strip(),contact.strip(),limit,opening,notes.strip()))
                    st.success("تم حفظ العميل.")
    with t2:
        rows=fetchall("SELECT * FROM customers ORDER BY active DESC,name")
        search=st.text_input("🔎 بحث",key="cust_search")
        data=[]
        for r in rows:
            if search and search.lower() not in (str(r["name"])+" "+str(r["phone"])).lower(): continue
            data.append(dict(r))
        st.dataframe(data,use_container_width=True,hide_index=True)

# =========================================================
# PRODUCTS & INVENTORY
# =========================================================
elif page == "المواد والمخزون":
    st.title("📦 المواد والمخزون")
    t1,t2,t3=st.tabs(["➕ مادة جديدة","📋 المواد","🔄 حركة المخزون"])
    with t1:
        with st.form("prod_add",clear_on_submit=True):
            a,b,c=st.columns(3)
            sku=a.text_input("SKU")
            name=b.text_input("اسم المادة *")
            brand=c.text_input("Brand")
            a,b,c=st.columns(3)
            cat=a.text_input("Category")
            unit=b.text_input("Unit",value="PCS")
            opening=c.number_input("Opening Qty",0.0,step=1.0)
            a,b,c=st.columns(3)
            cost=a.number_input("Cost",0.0,step=0.01)
            sale=b.number_input("Sale Price",0.0,step=0.01)
            minq=c.number_input("Min Qty",0.0,step=1.0)
            notes=st.text_area("Notes")
            if st.form_submit_button("حفظ المادة",type="primary"):
                if name.strip():
                    pid=execute("""INSERT INTO products
                        (sku,name,brand,category,unit,cost,sale_price,opening_qty,min_qty,notes)
                        VALUES (?,?,?,?,?,?,?,?,?,?)""",(sku.strip(),name.strip(),brand.strip(),cat.strip(),unit.strip(),cost,sale,opening,minq,notes.strip()))
                    if opening:
                        execute("""INSERT INTO inventory_movements
                            (movement_date,product_id,movement_type,qty,unit_cost,reference_type,reference_id,warehouse,notes)
                            VALUES (?,?,'Opening',?,?, 'Opening',?, 'Main','Opening balance')""",
                            (str(date.today()),pid,opening,cost,pid))
                    st.success("تم حفظ المادة.")
    with t2:
        rows=fetchall("""SELECT p.*,
            COALESCE((SELECT SUM(CASE
                WHEN movement_type IN ('Opening','Purchase','IN','Adjustment +') THEN qty
                ELSE -qty END)
                FROM inventory_movements m WHERE m.product_id=p.id),0) AS stock_qty
            FROM products p ORDER BY p.name""")
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)
    with t3:
        prods=fetchall("SELECT id,name,sku FROM products WHERE COALESCE(active,1)=1 ORDER BY name")
        if prods:
            pmap={f"{r['name']} — {r['sku'] or '-'}":r["id"] for r in prods}
            with st.form("inv_move",clear_on_submit=True):
                a,b,c=st.columns(3)
                plabel=a.selectbox("المادة",list(pmap.keys()))
                mtype=b.selectbox("نوع الحركة",["IN","OUT","Adjustment +","Adjustment -"])
                qty=c.number_input("Qty",0.0,step=1.0)
                a,b=st.columns(2)
                warehouse=a.text_input("Warehouse",value="Main")
                note=b.text_input("Note")
                if st.form_submit_button("حفظ الحركة",type="primary"):
                    execute("""INSERT INTO inventory_movements
                        (movement_date,product_id,movement_type,qty,unit_cost,reference_type,warehouse,notes)
                        VALUES (?,?,?,?,0,'Manual',?,?)""",
                        (str(date.today()),pmap[plabel],mtype,qty,warehouse.strip(),note.strip()))
                    st.success("تم.")

# =========================================================
# SALES
# =========================================================
elif page == "المبيعات":
    st.title("🧾 المبيعات")
    customers=fetchall("SELECT id,name FROM customers WHERE COALESCE(active,1)=1 ORDER BY name")
    cmap={"Cash / بدون عميل":None}
    for r in customers: cmap[r["name"]]=r["id"]
    with st.form("sale_head",clear_on_submit=True):
        a,b,c=st.columns(3)
        inv=a.text_input("Invoice No.")
        sdate=b.date_input("Date",date.today())
        clabel=c.selectbox("Customer",list(cmap.keys()))
        a,b,c=st.columns(3)
        total=a.number_input("Total Amount",0.0,step=100.0)
        paid=b.number_input("Paid Amount",0.0,step=100.0)
        curr=c.selectbox("Currency",["USD","IQD","CNY","EUR"])
        note=st.text_area("Notes")
        if st.form_submit_button("حفظ الفاتورة",type="primary"):
            sid=execute("""INSERT INTO sales
                (invoice_no,sale_date,customer_id,total_amount,paid_amount,currency,status,notes)
                VALUES (?,?,?,?,?,?,'Open',?)""",
                (inv.strip(),str(sdate),cmap[clabel],total,paid,curr,note.strip()))
            st.success(f"تم حفظ الفاتورة #{sid}")
    rows=fetchall("""SELECT s.*,c.name customer_name FROM sales s
        LEFT JOIN customers c ON c.id=s.customer_id ORDER BY s.id DESC""")
    if rows:
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)

# =========================================================
# PURCHASES
# =========================================================
elif page == "المشتريات":
    st.title("🛒 المشتريات")
    fmap=factory_map(False)
    suppliers=fetchall("SELECT id,name FROM suppliers WHERE COALESCE(active,1)=1 ORDER BY name")
    smap={"بدون مجهز":None}
    for r in suppliers: smap[r["name"]]=r["id"]
    with st.form("purchase_head",clear_on_submit=True):
        a,b,c=st.columns(3)
        pno=a.text_input("Purchase No.")
        pdate=b.date_input("Date",date.today())
        slabel=c.selectbox("Supplier",list(smap.keys()))
        flabel=st.selectbox("Factory (optional)",["بدون معمل"]+list(fmap.keys()))
        a,b,c=st.columns(3)
        total=a.number_input("Total",0.0,step=100.0)
        paid=b.number_input("Paid",0.0,step=100.0)
        curr=c.selectbox("Currency",["USD","IQD","CNY","EUR"])
        note=st.text_area("Notes")
        if st.form_submit_button("حفظ المشتريات",type="primary"):
            fid=None if flabel=="بدون معمل" else fmap[flabel]
            execute("""INSERT INTO purchases
                (purchase_no,purchase_date,supplier_id,factory_id,total_amount,paid_amount,currency,status,notes)
                VALUES (?,?,?,?,?,?,?,'Open',?)""",
                (pno.strip(),str(pdate),smap[slabel],fid,total,paid,curr,note.strip()))
            st.success("تم الحفظ.")
    rows=fetchall("""SELECT p.*,s.name supplier_name,f.factory_name FROM purchases p
        LEFT JOIN suppliers s ON s.id=p.supplier_id
        LEFT JOIN factories f ON f.id=p.factory_id ORDER BY p.id DESC""")
    if rows:
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)

# =========================================================
# CASH
# =========================================================
elif page == "الصندوق":
    st.title("💵 الصندوق")
    with st.form("cash_tx",clear_on_submit=True):
        a,b,c=st.columns(3)
        dt=a.date_input("Date",date.today())
        typ=b.selectbox("Type",["IN","OUT"])
        amount=c.number_input("Amount",0.0,step=100.0)
        a,b,c=st.columns(3)
        curr=a.selectbox("Currency",["USD","IQD","CNY","EUR"])
        party=b.text_input("Party")
        ref=c.text_input("Reference")
        note=st.text_area("Notes")
        if st.form_submit_button("حفظ الحركة",type="primary"):
            execute("""INSERT INTO cash_transactions
                (trans_date,trans_type,amount,currency,party,reference,notes)
                VALUES (?,?,?,?,?,?,?)""",
                (str(dt),typ,amount,curr,party.strip(),ref.strip(),note.strip()))
            st.success("تم.")
    rows=fetchall("SELECT * FROM cash_transactions ORDER BY id DESC")
    if rows:
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)

# =========================================================
# DOCUMENTS
# =========================================================
elif page == "المستندات":
    st.title("📁 المستندات")
    st.caption("عرض سريع لكل مستندات المعامل والطلبيات والمشحونات والدفعات.")
    tabs=st.tabs(["Factories","Orders","Shipments","Payments"])
    with tabs[0]:
        rows=fetchall("""SELECT d.*,f.factory_name FROM factory_documents d
            LEFT JOIN factories f ON f.id=d.factory_id ORDER BY d.id DESC""")
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)
    with tabs[1]:
        rows=fetchall("""SELECT d.*,o.pi_no FROM order_documents d
            LEFT JOIN orders o ON o.id=d.order_id ORDER BY d.id DESC""")
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)
    with tabs[2]:
        rows=fetchall("""SELECT d.*,s.bl_no FROM shipment_documents d
            LEFT JOIN shipments s ON s.id=d.shipment_id ORDER BY d.id DESC""")
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)
    with tabs[3]:
        rows=fetchall("""SELECT d.*,p.reference_no FROM payment_documents d
            LEFT JOIN payments p ON p.id=d.payment_id ORDER BY d.id DESC""")
        st.dataframe([dict(r) for r in rows],use_container_width=True,hide_index=True)

# =========================================================
# LEGACY ACCESS
# =========================================================
elif page == "نظام Access القديم":
    st.title("🗃️ نظام Access القديم")
    if not access_db_available():
        st.error("ارفع mido_access_structure_only.db بجانب app.py في GitHub.")
    else:
        tables=access_table_names()
        a,b,c=st.columns(3)
        a.metric("Tables",len(tables))
        b.metric("Access Structure","Linked ✅")
        c.metric("Old Data","Not included")
        st.info("هذه نسخة هيكل Access فقط. البيانات القديمة غير موجودة في GitHub عمداً حتى يبقى المشروع خفيفاً.")

        selected=st.selectbox("اختر جدول",tables)
        cols=access_table_columns(selected)
        count=access_table_count(selected)
        a,b=st.columns(2)
        a.metric("Fields",len(cols))
        b.metric("Rows",count)
        st.code("\n".join(cols),language=None)

# =========================================================
# REPORTS
# =========================================================
elif page == "التقارير":
    st.title("📄 التقارير")
    st.subheader("Shipment Report")
    rows=fetchall("""SELECT s.*,f.factory_name,o.pi_no FROM shipments s
        LEFT JOIN factories f ON f.id=s.factory_id
        LEFT JOIN orders o ON o.id=s.order_id ORDER BY s.id DESC""")
    table=[]
    for s in rows:
        table.append({
            "PI No.":s["pi_no"] or "-","Factory":s["factory_name"] or "-",
            "BL":s["bl_no"] or "-","Ship Com.":s["shipping_company"] or "-",
            "Destination":s["destination"] or "-","CONTS":s["containers_count"] or 0,
            "Type":s["container_type"] or "40HQ","CI Amount/USD":s["ci_amount"] or 0,
            "Shipping":s["shipping_cost"] or 0,"Total Amount":s["total_amount"] or 0,
            "ETD":s["etd"] or "-","ETA":s["eta"] or "-",
            "State":s["shipment_status"] or "-","Payment":s["payment_status"] or "-"
        })
    if table:
        st.dataframe(table,use_container_width=True,hide_index=True)
        st.download_button("📄 Download Shipment PDF",shipments_pdf(table),
            file_name=f"MIDO_Main_Shipments_{date.today()}.pdf",mime="application/pdf")
        csv_bytes=pd.DataFrame(table).to_csv(index=False).encode("utf-8-sig")
        st.download_button("📊 Download Excel-compatible CSV",csv_bytes,
            file_name=f"MIDO_Main_Shipments_{date.today()}.csv",mime="text/csv")
    else:
        st.info("لا توجد بيانات للتقرير.")

st.sidebar.divider()
st.sidebar.caption("MIDO Consolidated Build — AI intentionally excluded for now")
