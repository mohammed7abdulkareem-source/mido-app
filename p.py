import streamlit as st
import sqlite3
from pathlib import Path
from datetime import datetime, date
import re
import shutil

DB_FILE = Path("mido_v2.db")
ACCESS_DB_FILE = Path("mido_access_structure_only.db")
UPLOAD_ROOT = Path("uploads")
FACTORY_UPLOAD_ROOT = UPLOAD_ROOT / "factories"
ORDER_UPLOAD_ROOT = UPLOAD_ROOT / "orders"
SHIPMENT_UPLOAD_ROOT = UPLOAD_ROOT / "shipments"
SHIPMENT_INBOX_ROOT = UPLOAD_ROOT / "shipment_inbox"

st.set_page_config(
    page_title="MIDO",
    page_icon="📦",
    layout="wide",
)

# =========================================================
# DATABASE HELPERS
# =========================================================
def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def table_exists(table_name):
    conn = get_conn()
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table_name,)
    ).fetchone()
    conn.close()
    return row is not None

def get_columns(table_name):
    conn = get_conn()
    rows = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
    conn.close()
    return {r["name"] for r in rows}

def add_missing_columns(table_name, definitions):
    existing = get_columns(table_name)
    conn = get_conn()
    for definition in definitions:
        column_name = definition.split()[0]
        if column_name not in existing:
            conn.execute(f"ALTER TABLE {table_name} ADD COLUMN {definition}")
    conn.commit()
    conn.close()

def init_database():
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS app_settings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            setting_key TEXT UNIQUE NOT NULL,
            setting_value TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS factories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factory_name TEXT NOT NULL,
            chinese_name TEXT,
            brand_name TEXT,
            contact_person TEXT,
            phone TEXT,
            wechat TEXT,
            email TEXT,
            address TEXT,
            website TEXT,
            bank_info TEXT,
            payment_terms TEXT,
            monthly_order_deadline INTEGER DEFAULT 20,
            private_notes TEXT,
            public_notes TEXT,
            active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS factory_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factory_id INTEGER NOT NULL,
            document_type TEXT,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factory_id INTEGER NOT NULL
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS order_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            order_id INTEGER NOT NULL,
            size TEXT,
            pattern TEXT,
            description TEXT,
            quantity INTEGER DEFAULT 0,
            unit_price REAL DEFAULT 0,
            amount REAL DEFAULT 0,
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS order_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            order_id INTEGER NOT NULL,
            document_type TEXT,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS shipments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factory_id INTEGER NOT NULL,
            order_id INTEGER,
            shipment_no TEXT,
            bl_no TEXT,
            shipping_company TEXT,
            destination TEXT,
            container_type TEXT DEFAULT '40HQ',
            containers_count REAL DEFAULT 0,
            ci_amount REAL DEFAULT 0,
            shipping_cost REAL DEFAULT 0,
            total_amount REAL DEFAULT 0,
            etd TEXT,
            eta TEXT,
            shipment_status TEXT DEFAULT 'Preparing',
            payment_status TEXT DEFAULT 'Not Paid',
            private_notes TEXT,
            public_notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS shipment_containers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            shipment_id INTEGER NOT NULL,
            container_no TEXT,
            seal_no TEXT,
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS shipment_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            shipment_id INTEGER NOT NULL,
            document_type TEXT,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS shipment_intakes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factory_id INTEGER NOT NULL,
            order_id INTEGER,
            intake_name TEXT,
            status TEXT DEFAULT 'Uploaded',
            private_notes TEXT,
            extracted_pi_no TEXT,
            extracted_bl_no TEXT,
            extracted_shipping_company TEXT,
            extracted_destination TEXT,
            extracted_container_type TEXT,
            extracted_containers_count REAL DEFAULT 0,
            extracted_ci_amount REAL DEFAULT 0,
            extracted_shipping_cost REAL DEFAULT 0,
            extracted_etd TEXT,
            extracted_eta TEXT,
            extracted_container_numbers TEXT,
            analysis_notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS shipment_intake_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id INTEGER NOT NULL,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_ext TEXT,
            detected_type TEXT DEFAULT 'Unclassified',
            uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    conn.commit()
    conn.close()

    # IMPORTANT: migrate old database safely instead of deleting it
    add_missing_columns("factories", [
        "chinese_name TEXT",
        "brand_name TEXT",
        "contact_person TEXT",
        "phone TEXT",
        "wechat TEXT",
        "email TEXT",
        "address TEXT",
        "website TEXT",
        "bank_info TEXT",
        "payment_terms TEXT",
        "monthly_order_deadline INTEGER DEFAULT 20",
        "private_notes TEXT",
        "public_notes TEXT",
        "active INTEGER DEFAULT 1",
        "created_at TEXT",
        "updated_at TEXT",
    ])

    add_missing_columns("orders", [
        "factory_id INTEGER",
        "pi_no TEXT",
        "order_date TEXT",
        "production_due_date TEXT",
        "order_status TEXT DEFAULT 'Draft'",
        "currency TEXT DEFAULT 'USD'",
        "pi_amount REAL DEFAULT 0",
        "expected_containers REAL DEFAULT 0",
        "ordered_containers REAL DEFAULT 0",
        "remaining_containers REAL DEFAULT 0",
        "shipping_estimate REAL DEFAULT 0",
        "total_estimated_amount REAL DEFAULT 0",
        "destination TEXT",
        "private_notes TEXT",
        "public_notes TEXT",
        "created_at TEXT",
        "updated_at TEXT",
    ])

    add_missing_columns("factory_documents", [
        "factory_id INTEGER",
        "document_type TEXT",
        "file_name TEXT",
        "file_path TEXT",
        "uploaded_at TEXT",
    ])

    add_missing_columns("order_items", [
        "order_id INTEGER",
        "size TEXT",
        "pattern TEXT",
        "description TEXT",
        "quantity INTEGER DEFAULT 0",
        "unit_price REAL DEFAULT 0",
        "amount REAL DEFAULT 0",
        "notes TEXT",
        "created_at TEXT",
    ])

    add_missing_columns("order_documents", [
        "order_id INTEGER",
        "document_type TEXT",
        "file_name TEXT",
        "file_path TEXT",
        "uploaded_at TEXT",
    ])

    add_missing_columns("shipments", [
        "factory_id INTEGER",
        "order_id INTEGER",
        "shipment_no TEXT",
        "bl_no TEXT",
        "shipping_company TEXT",
        "destination TEXT",
        "container_type TEXT DEFAULT '40HQ'",
        "containers_count REAL DEFAULT 0",
        "ci_amount REAL DEFAULT 0",
        "shipping_cost REAL DEFAULT 0",
        "total_amount REAL DEFAULT 0",
        "etd TEXT",
        "eta TEXT",
        "shipment_status TEXT DEFAULT 'Preparing'",
        "payment_status TEXT DEFAULT 'Not Paid'",
        "private_notes TEXT",
        "public_notes TEXT",
        "created_at TEXT",
        "updated_at TEXT",
    ])

    add_missing_columns("shipment_containers", [
        "shipment_id INTEGER",
        "container_no TEXT",
        "seal_no TEXT",
        "notes TEXT",
        "created_at TEXT",
    ])

    add_missing_columns("shipment_documents", [
        "shipment_id INTEGER",
        "document_type TEXT",
        "file_name TEXT",
        "file_path TEXT",
        "uploaded_at TEXT",
    ])

    add_missing_columns("shipment_intakes", [
        "factory_id INTEGER",
        "order_id INTEGER",
        "intake_name TEXT",
        "status TEXT DEFAULT 'Uploaded'",
        "private_notes TEXT",
        "extracted_pi_no TEXT",
        "extracted_bl_no TEXT",
        "extracted_shipping_company TEXT",
        "extracted_destination TEXT",
        "extracted_container_type TEXT",
        "extracted_containers_count REAL DEFAULT 0",
        "extracted_ci_amount REAL DEFAULT 0",
        "extracted_shipping_cost REAL DEFAULT 0",
        "extracted_etd TEXT",
        "extracted_eta TEXT",
        "extracted_container_numbers TEXT",
        "analysis_notes TEXT",
        "created_at TEXT",
    ])

    add_missing_columns("shipment_intake_files", [
        "intake_id INTEGER",
        "file_name TEXT",
        "file_path TEXT",
        "file_ext TEXT",
        "detected_type TEXT DEFAULT 'Unclassified'",
        "uploaded_at TEXT",
    ])

init_database()

def fetchall(query, params=()):
    conn = get_conn()
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return rows

def fetchone(query, params=()):
    conn = get_conn()
    row = conn.execute(query, params).fetchone()
    conn.close()
    return row

def execute(query, params=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(query, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id

def save_uploaded_files(uploaded_files, target_dir, table_name, entity_field, entity_id, document_type):
    target_dir.mkdir(parents=True, exist_ok=True)
    saved = 0

    for uploaded in uploaded_files:
        filename = Path(uploaded.name).name
        target = target_dir / filename

        if target.exists():
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            target = target_dir / f"{target.stem}_{stamp}{target.suffix}"

        target.write_bytes(uploaded.getbuffer())

        execute(
            f"""
            INSERT INTO {table_name}
            ({entity_field}, document_type, file_name, file_path)
            VALUES (?, ?, ?, ?)
            """,
            (entity_id, document_type, target.name, str(target)),
        )
        saved += 1

    return saved


# =========================================================
# FILE ANALYSIS (NO AI)
# =========================================================
def normalize_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()

def read_file_text(path):
    path = Path(path)
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        elif ext in (".xlsx", ".xlsm"):
            from openpyxl import load_workbook
            wb = load_workbook(str(path), data_only=True, read_only=True)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(f"SHEET {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(v) for v in row if v is not None]
                    if vals:
                        chunks.append(" | ".join(vals))
            return "\n".join(chunks)
        elif ext == ".csv":
            import csv
            chunks = []
            with open(path, "r", encoding="utf-8-sig", errors="ignore", newline="") as f:
                for row in csv.reader(f):
                    chunks.append(" | ".join(row))
            return "\n".join(chunks)
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(path))
            chunks = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    chunks.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(chunks)
        elif ext in (".txt",):
            return path.read_text(encoding="utf-8", errors="ignore")
        else:
            return ""
    except Exception:
        return ""

def detect_document_type(file_name, text):
    s = (file_name + " " + text[:4000]).lower()
    if "packing list" in s or "packing" in s:
        return "Packing List"
    if "commercial invoice" in s or ("invoice" in s and "proforma" not in s):
        return "Commercial Invoice"
    if "certificate of origin" in s or "country of origin" in s:
        return "Certificate of Origin (CO)"
    if re.search(r"\bcoc\b", s) or "certificate of conformity" in s:
        return "COC"
    if "bill of lading" in s or re.search(r"\bb/?l\b", s):
        return "Bill of Lading"
    if "proforma" in s or re.search(r"\bpi\b", s):
        return "PI / Proforma Invoice"
    return "Unclassified"

def find_labeled_value(text, labels, token_pattern=r"[^|\n]{1,80}"):
    for label in labels:
        m = re.search(rf"{label}\s*[:#\.\-]*\s*({token_pattern})", text, flags=re.I)
        if m:
            return normalize_text(m.group(1)).strip(" :|-_")
    return ""

def find_amount_after_labels(text, labels):
    for label in labels:
        m = re.search(rf"{label}[^\d$]{{0,30}}\$?\s*([\d,]+(?:\.\d+)?)", text, flags=re.I)
        if m:
            try:
                return float(m.group(1).replace(",", ""))
            except Exception:
                pass
    return 0.0

def find_date_after_label(text, label):
    m = re.search(
        rf"{label}[^0-9]{{0,25}}((?:20\d{{2}}[./-]\d{{1,2}}[./-]\d{{1,2}})|(?:\d{{1,2}}[./-]\d{{1,2}}[./-]20\d{{2}}))",
        text, flags=re.I
    )
    return m.group(1).replace(".", "-").replace("/", "-") if m else ""

def analyze_intake_files(intake_id):
    file_rows = fetchall("""
        SELECT * FROM shipment_intake_files
        WHERE intake_id=? ORDER BY id
    """, (intake_id,))

    all_text_parts = []
    all_names = []
    notes = []
    for f in file_rows:
        path = Path(f["file_path"])
        text = read_file_text(path)
        all_names.append(f["file_name"] or "")
        all_text_parts.append(text)
        detected = detect_document_type(f["file_name"] or "", text)
        execute("UPDATE shipment_intake_files SET detected_type=? WHERE id=?", (detected, f["id"]))
        if not text and path.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
            notes.append(f"الصورة {f['file_name']} محفوظة، لكن قراءة محتوى الصور ستتم لاحقاً مع محرك AI.")
        elif not text and path.suffix.lower() in (".xls", ".doc"):
            notes.append(f"الملف القديم {f['file_name']} محفوظ، لكن التحليل المباشر يفضّل XLSX/DOCX.")

    full_text = "\n".join(all_text_parts)
    searchable = full_text + "\n" + "\n".join(all_names)

    pi_no = find_labeled_value(searchable, [r"PI\s*(?:NO|NUMBER)", r"PROFORMA\s+INVOICE\s*(?:NO|NUMBER)"], r"[A-Z0-9][A-Z0-9._/-]{2,40}")
    bl_no = find_labeled_value(searchable, [r"B/?L\s*(?:NO|NUMBER)?", r"BILL\s+OF\s+LADING\s*(?:NO|NUMBER)?"], r"[A-Z0-9][A-Z0-9._/-]{5,40}")

    shipping_company = ""
    shipping_map = ["MSC", "HPL", "CMA", "OOCL", "WANHAI", "ASYAD", "PIL", "MAERSK", "COSCO", "EVERGREEN"]
    upper = searchable.upper()
    for company in shipping_map:
        if re.search(rf"\b{re.escape(company)}\b", upper):
            shipping_company = company
            break

    destination = ""
    destinations = [
        ("UMM QASR", ["UMM QASR", "UMM,IRAQ", "UMM IRAQ"]),
        ("AQABA", ["AQABA", "AQABA,JOR", "AQABA JOR"]),
        ("MERSIN", ["MERSIN", "MERSIN,TUR", "MERSIN TUR"]),
    ]
    for canonical, variants in destinations:
        if any(v in upper for v in variants):
            destination = canonical
            break

    container_numbers = sorted(set(re.findall(r"\b[A-Z]{4}\d{7}\b", upper)))

    containers_count = 0.0
    count_patterns = [
        r"(\d+(?:\.\d+)?)\s*[*X×]\s*40HQ",
        r"CONTS?\s*[*X×]?\s*40HQ[^\d]{0,20}(\d+(?:\.\d+)?)",
        r"(\d+(?:\.\d+)?)\s*(?:CONTAINER|CONTAINERS|CNTR|CNTRS)\b",
    ]
    for pat in count_patterns:
        m = re.search(pat, upper, flags=re.I)
        if m:
            try:
                containers_count = float(m.group(1))
                break
            except Exception:
                pass
    if not containers_count and container_numbers:
        containers_count = float(len(container_numbers))

    container_type = "40HQ" if "40HQ" in upper else ("40GP" if "40GP" in upper else ("20GP" if "20GP" in upper else "40HQ"))
    ci_amount = find_amount_after_labels(searchable, [r"CI\s*AMOUNT(?:/USD)?", r"COMMERCIAL\s+INVOICE\s+AMOUNT", r"TOTAL\s+INVOICE\s+AMOUNT"])
    shipping_cost = find_amount_after_labels(searchable, [r"SHIPPING(?:\s+COST)?", r"FREIGHT"])
    etd = find_date_after_label(searchable, r"ETD")
    eta = find_date_after_label(searchable, r"ETA")

    execute("""
        UPDATE shipment_intakes
        SET extracted_pi_no=?, extracted_bl_no=?, extracted_shipping_company=?,
            extracted_destination=?, extracted_container_type=?, extracted_containers_count=?,
            extracted_ci_amount=?, extracted_shipping_cost=?, extracted_etd=?, extracted_eta=?,
            extracted_container_numbers=?, analysis_notes=?, status='Analyzed'
        WHERE id=?
    """, (
        pi_no, bl_no, shipping_company, destination, container_type,
        containers_count, ci_amount, shipping_cost, etd, eta,
        "\n".join(container_numbers), "\n".join(notes), intake_id
    ))

    return fetchone("SELECT * FROM shipment_intakes WHERE id=?", (intake_id,))


# =========================================================
# ACCESS LEGACY DATABASE
# =========================================================
def access_db_available():
    return ACCESS_DB_FILE.exists()

def get_access_conn():
    conn = sqlite3.connect(ACCESS_DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def access_table_names():
    if not access_db_available():
        return []
    conn = get_access_conn()
    rows = conn.execute("""
        SELECT name
        FROM sqlite_master
        WHERE type='table'
          AND name NOT LIKE 'sqlite_%'
          AND name NOT LIKE '\\_%' ESCAPE '\\'
        ORDER BY name
    """).fetchall()
    conn.close()
    return [r["name"] for r in rows]

def access_table_columns(table_name):
    conn = get_access_conn()
    rows = conn.execute(f'PRAGMA table_info("{table_name}")').fetchall()
    conn.close()
    return [r["name"] for r in rows]

def access_table_count(table_name):
    conn = get_access_conn()
    try:
        n = conn.execute(f'SELECT COUNT(*) AS n FROM "{table_name}"').fetchone()["n"]
    finally:
        conn.close()
    return n

def access_table_rows(table_name, limit=200):
    conn = get_access_conn()
    try:
        rows = conn.execute(f'SELECT * FROM "{table_name}" LIMIT ?', (int(limit),)).fetchall()
    finally:
        conn.close()
    return rows

# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.title("MIDO")
st.sidebar.caption("Business & Shipment Management System")

page = st.sidebar.radio(
    "القائمة الرئيسية",
    ["الرئيسية", "المعامل", "الطلبيات", "المشحونات", "الدفعات", "المستندات", "نظام Access القديم", "التقارير"],
)

# =========================================================
# HOME
# =========================================================
if page == "الرئيسية":
    st.title("📦 MIDO")
    st.subheader("نظام إدارة المعامل والطلبيات والمشحونات")

    factory_count = fetchone("SELECT COUNT(*) AS n FROM factories")["n"]
    order_count = fetchone("SELECT COUNT(*) AS n FROM orders")["n"]
    shipment_count = fetchone("SELECT COUNT(*) AS n FROM shipments")["n"]
    in_transit_containers = fetchone("""
        SELECT COALESCE(SUM(containers_count),0) AS n
        FROM shipments
        WHERE shipment_status IN ('Booked','Shipped','In Transit')
    """)["n"]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("المعامل", factory_count)
    c2.metric("الطلبيات", order_count)
    c3.metric("المشحونات", shipment_count)
    c4.metric("الحاويات بالطريق", f"{in_transit_containers:g}")

    st.success("قاعدة البيانات تم تحديثها تلقائياً بدون حذف بياناتك.")

    st.subheader("آخر الطلبيات")
    recent_orders = fetchall("""
        SELECT o.id, o.pi_no, f.factory_name, o.order_status,
               o.pi_amount, o.currency, o.expected_containers,
               o.production_due_date
        FROM orders o
        LEFT JOIN factories f ON f.id = o.factory_id
        ORDER BY o.id DESC
        LIMIT 10
    """)

    if recent_orders:
        st.dataframe(
            [
                {
                    "ID": r["id"],
                    "PI No.": r["pi_no"] or "-",
                    "المعمل": r["factory_name"] or "-",
                    "الحالة": r["order_status"] or "-",
                    "PI Amount": r["pi_amount"] or 0,
                    "Currency": r["currency"] or "USD",
                    "الحاويات المطلوبة": r["expected_containers"] or 0,
                    "موعد الإنتاج": r["production_due_date"] or "-",
                }
                for r in recent_orders
            ],
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.info("لا توجد طلبيات بعد.")

# =========================================================
# FACTORIES
# =========================================================
elif page == "المعامل":
    st.title("🏭 المعامل الصينية")
    tab1, tab2 = st.tabs(["➕ إضافة معمل", "📋 المعامل المسجلة"])

    with tab1:
        with st.form("add_factory_form", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)
            factory_name = c1.text_input("اسم المعمل *")
            chinese_name = c2.text_input("اسم المعمل بالصيني")
            brand_name = c3.text_input("العلامة / البراند")

            contact_person = c1.text_input("الشخص المسؤول")
            phone = c2.text_input("رقم الهاتف")
            wechat = c3.text_input("WeChat")

            email = c1.text_input("Email")
            website = c2.text_input("Website")
            monthly_order_deadline = c3.number_input(
                "موعد تجهيز الطلبية الشهري",
                min_value=1,
                max_value=31,
                value=20,
                step=1,
            )

            address = st.text_area("عنوان المعمل")
            bank_info = st.text_area("معلومات الحساب البنكي")
            payment_terms = st.text_area("شروط الدفع")

            n1, n2 = st.columns(2)
            private_notes = n1.text_area("🔒 ملاحظات خاصة")
            public_notes = n2.text_area("📝 ملاحظات عامة")

            submitted = st.form_submit_button("💾 حفظ المعمل", type="primary")

            if submitted:
                if not factory_name.strip():
                    st.error("يجب كتابة اسم المعمل.")
                else:
                    execute("""
                        INSERT INTO factories (
                            factory_name, chinese_name, brand_name,
                            contact_person, phone, wechat, email,
                            address, website, bank_info, payment_terms,
                            monthly_order_deadline, private_notes, public_notes,
                            active, created_at, updated_at
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """, (
                        factory_name.strip(), chinese_name.strip(), brand_name.strip(),
                        contact_person.strip(), phone.strip(), wechat.strip(), email.strip(),
                        address.strip(), website.strip(), bank_info.strip(), payment_terms.strip(),
                        int(monthly_order_deadline), private_notes.strip(), public_notes.strip(),
                    ))
                    st.success("تم حفظ المعمل.")

    with tab2:
        factories = fetchall("SELECT * FROM factories ORDER BY active DESC, factory_name")
        if not factories:
            st.info("لا توجد معامل بعد.")
        else:
            for factory in factories:
                title = f"{'🟢' if factory['active'] else '⚪'} {factory['factory_name']}"
                if factory["brand_name"]:
                    title += f" — {factory['brand_name']}"

                with st.expander(title):
                    c1, c2, c3 = st.columns(3)
                    c1.write(f"**الشخص المسؤول:** {factory['contact_person'] or '-'}")
                    c1.write(f"**الهاتف:** {factory['phone'] or '-'}")
                    c1.write(f"**WeChat:** {factory['wechat'] or '-'}")
                    c2.write(f"**Email:** {factory['email'] or '-'}")
                    c2.write(f"**Website:** {factory['website'] or '-'}")
                    c2.write(f"**موعد الطلب:** قبل يوم {factory['monthly_order_deadline'] or 20}")
                    c3.write(f"**البراند:** {factory['brand_name'] or '-'}")
                    c3.write(f"**رقم المعمل:** {factory['id']}")

                    n1, n2 = st.columns(2)
                    n1.info(factory["private_notes"] or "لا توجد ملاحظات خاصة.")
                    n2.info(factory["public_notes"] or "لا توجد ملاحظات عامة.")

                    st.markdown("### 📁 ملفات المعمل")
                    doc_type = st.selectbox(
                        "نوع الملف",
                        [
                            "Company Registration",
                            "Contract / Agreement",
                            "Agency / Authorization",
                            "Bank Details",
                            "Price List",
                            "Certificate",
                            "PI / Quotation",
                            "Other",
                        ],
                        key=f"factory_doc_type_{factory['id']}",
                    )
                    uploads = st.file_uploader(
                        "ارفع ملف أو أكثر",
                        accept_multiple_files=True,
                        type=["pdf", "xlsx", "xls", "docx", "doc", "png", "jpg", "jpeg", "webp"],
                        key=f"factory_upload_{factory['id']}",
                    )

                    if uploads and st.button("⬆️ حفظ الملفات", key=f"save_factory_{factory['id']}"):
                        saved = save_uploaded_files(
                            uploads,
                            FACTORY_UPLOAD_ROOT / str(factory["id"]),
                            "factory_documents",
                            "factory_id",
                            factory["id"],
                            doc_type,
                        )
                        st.success(f"تم حفظ {saved} ملف/ملفات.")

# =========================================================
# ORDERS
# =========================================================
elif page == "الطلبيات":
    st.title("📋 الطلبيات")

    factories = fetchall("""
        SELECT id, factory_name, brand_name
        FROM factories
        WHERE COALESCE(active,1)=1
        ORDER BY factory_name
    """)

    if not factories:
        st.warning("أضف معمل أولاً.")
    else:
        factory_map = {
            f"{r['factory_name']} — {r['brand_name'] or 'بدون براند'}": r["id"]
            for r in factories
        }

        tab1, tab2 = st.tabs(["➕ طلبية جديدة", "📋 الطلبيات المسجلة"])

        with tab1:
            with st.form("new_order_form", clear_on_submit=True):
                c1, c2, c3 = st.columns(3)
                factory_label = c1.selectbox("المعمل *", list(factory_map.keys()))
                pi_no = c2.text_input("PI No. / رقم الطلبية *")
                order_status = c3.selectbox(
                    "حالة الطلبية",
                    ["Draft", "Confirmed", "In Production", "Ready", "Partially Shipped", "Completed", "Cancelled"]
                )

                c4, c5, c6 = st.columns(3)
                order_date = c4.date_input("تاريخ الطلبية", value=date.today())
                production_due_date = c5.date_input("موعد الإنتاج / الجاهزية", value=date.today())
                currency = c6.selectbox("العملة", ["USD", "CNY", "EUR"])

                c7, c8, c9 = st.columns(3)
                pi_amount = c7.number_input("PI Amount", min_value=0.0, step=100.0, format="%.2f")
                expected_containers = c8.number_input("عدد الحاويات المطلوبة", min_value=0.0, step=0.5)
                ordered_containers = c9.number_input("عدد الحاويات المؤكدة", min_value=0.0, step=0.5)

                c10, c11, c12 = st.columns(3)
                shipping_estimate = c10.number_input("تكلفة الشحن التقديرية", min_value=0.0, step=100.0)
                destination = c11.text_input("الوجهة")
                total_estimated_amount = c12.number_input("الإجمالي التقديري", min_value=0.0, step=100.0)

                n1, n2 = st.columns(2)
                private_notes = n1.text_area("🔒 ملاحظات خاصة")
                public_notes = n2.text_area("📝 ملاحظات عامة")

                submit_order = st.form_submit_button("💾 حفظ الطلبية", type="primary")

                if submit_order:
                    if not pi_no.strip():
                        st.error("يجب كتابة PI No.")
                    else:
                        remaining = max(float(expected_containers) - float(ordered_containers), 0)
                        order_id = execute("""
                            INSERT INTO orders (
                                factory_id, pi_no, order_date, production_due_date,
                                order_status, currency, pi_amount,
                                expected_containers, ordered_containers,
                                remaining_containers, shipping_estimate,
                                total_estimated_amount, destination,
                                private_notes, public_notes, created_at, updated_at
                            )
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                        """, (
                            factory_map[factory_label], pi_no.strip(), str(order_date),
                            str(production_due_date), order_status, currency, float(pi_amount),
                            float(expected_containers), float(ordered_containers), remaining,
                            float(shipping_estimate), float(total_estimated_amount),
                            destination.strip(), private_notes.strip(), public_notes.strip(),
                        ))
                        st.success(f"تم حفظ الطلبية رقم {order_id}")

        with tab2:
            orders = fetchall("""
                SELECT o.*, f.factory_name, f.brand_name
                FROM orders o
                LEFT JOIN factories f ON f.id = o.factory_id
                ORDER BY o.id DESC
            """)

            if not orders:
                st.info("لا توجد طلبيات.")
            else:
                for order in orders:
                    header = f"#{order['id']} — {order['pi_no'] or 'بدون PI'} — {order['factory_name'] or 'بدون معمل'}"

                    with st.expander(header):
                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("PI Amount", f"{float(order['pi_amount'] or 0):,.2f} {order['currency'] or 'USD'}")
                        m2.metric("الحاويات المطلوبة", f"{float(order['expected_containers'] or 0):g}")
                        m3.metric("المؤكد", f"{float(order['ordered_containers'] or 0):g}")
                        m4.metric("المتبقي", f"{float(order['remaining_containers'] or 0):g}")

                        st.markdown("### 📦 تفاصيل الطلبية")

                        with st.form(f"item_form_{order['id']}", clear_on_submit=True):
                            i1, i2, i3 = st.columns(3)
                            size = i1.text_input("القياس / Size")
                            pattern = i2.text_input("النقشة / Pattern")
                            description = i3.text_input("الوصف")

                            i4, i5, i6 = st.columns(3)
                            quantity = i4.number_input("العدد / Qty", min_value=0, step=1)
                            unit_price = i5.number_input("سعر الوحدة", min_value=0.0, step=0.01)
                            item_notes = i6.text_input("ملاحظة")

                            if st.form_submit_button("➕ إضافة السطر"):
                                amount = int(quantity) * float(unit_price)
                                execute("""
                                    INSERT INTO order_items (
                                        order_id, size, pattern, description,
                                        quantity, unit_price, amount, notes, created_at
                                    )
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
                                """, (
                                    order["id"], size.strip(), pattern.strip(),
                                    description.strip(), int(quantity), float(unit_price),
                                    amount, item_notes.strip(),
                                ))
                                st.success("تمت إضافة السطر.")
                                st.rerun()

                        items = fetchall("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (order["id"],))
                        if items:
                            st.dataframe(
                                [
                                    {
                                        "Size": x["size"],
                                        "Pattern": x["pattern"],
                                        "Description": x["description"],
                                        "Qty": x["quantity"],
                                        "Unit Price": x["unit_price"],
                                        "Amount": x["amount"],
                                        "Notes": x["notes"],
                                    }
                                    for x in items
                                ],
                                use_container_width=True,
                                hide_index=True,
                            )

                        st.markdown("### 📁 ملفات الطلبية")
                        doc_type = st.selectbox(
                            "نوع الملف",
                            ["PI / Proforma Invoice", "Order Excel", "Order PDF", "Price List", "Specification", "Confirmation", "Other"],
                            key=f"order_doc_type_{order['id']}",
                        )
                        files = st.file_uploader(
                            "ارفع ملف أو أكثر",
                            accept_multiple_files=True,
                            type=["pdf", "xlsx", "xls", "csv", "docx", "doc", "png", "jpg", "jpeg", "webp"],
                            key=f"order_upload_{order['id']}",
                        )

                        if files and st.button("⬆️ حفظ ملفات الطلبية", key=f"save_order_{order['id']}"):
                            saved = save_uploaded_files(
                                files,
                                ORDER_UPLOAD_ROOT / str(order["id"]),
                                "order_documents",
                                "order_id",
                                order["id"],
                                doc_type,
                            )
                            st.success(f"تم حفظ {saved} ملف/ملفات.")

# =========================================================
# PLACEHOLDERS
# =========================================================
elif page == "المشحونات":
    st.title("🚢 المشحونات")

    factories = fetchall("""
        SELECT id, factory_name, brand_name
        FROM factories
        WHERE COALESCE(active,1)=1
        ORDER BY factory_name
    """)

    if not factories:
        st.warning("أضف معمل أولاً.")
    else:
        factory_map = {
            f"{r['factory_name']} — {r['brand_name'] or 'بدون براند'}": r["id"]
            for r in factories
        }

        tab1, tab2, tab3 = st.tabs(["📥 إضافة شحنة من الملفات", "✍️ إدخال يدوي", "📋 المشحونات الرئيسية"])

        # ---------------- FILE-FIRST SHIPMENT INTAKE ----------------
        with tab1:
            st.subheader("إضافة شحنة من الملفات")
            st.caption("اختَر المعمل وارفع كل ملفات الشحنة مرة واحدة. الملفات ستبقى مرتبطة بهذه الشحنة، ولن تحتاج لإعادة رفعها في قسم المستندات.")

            intake_factory_label = st.selectbox(
                "المعمل *",
                list(factory_map.keys()),
                key="intake_factory_select",
            )
            intake_factory_id = factory_map[intake_factory_label]

            intake_orders = fetchall("""
                SELECT id, pi_no
                FROM orders
                WHERE factory_id=?
                ORDER BY id DESC
            """, (intake_factory_id,))

            intake_order_map = {"بدون ربط بطلبية حالياً": None}
            for r in intake_orders:
                intake_order_map[f"{r['pi_no'] or 'Order'} — #{r['id']}"] = r["id"]

            intake_order_label = st.selectbox(
                "الطلبية المرتبطة (اختياري)",
                list(intake_order_map.keys()),
                key="intake_order_select",
            )

            intake_name = st.text_input(
                "اسم مؤقت للشحنة",
                placeholder="مثال: BOTRIAN AUG-2026 / Shipment 01",
                key="intake_name",
            )

            intake_files = st.file_uploader(
                "ارفع جميع ملفات الشحنة دفعة واحدة",
                accept_multiple_files=True,
                type=["pdf", "xlsx", "xls", "csv", "docx", "doc", "png", "jpg", "jpeg", "webp"],
                key="shipment_intake_upload",
                help="Commercial Invoice, Packing List, CO, COC, B/L, Excel, Word, صور وغيرها.",
            )

            intake_private_notes = st.text_area(
                "🔒 ملاحظات خاصة لهذه الشحنة",
                key="intake_private_notes",
            )

            if intake_files:
                st.markdown("#### الملفات الجاهزة للرفع")
                preview_rows = []
                for f in intake_files:
                    ext = Path(f.name).suffix.lower().replace(".", "")
                    preview_rows.append({
                        "اسم الملف": f.name,
                        "النوع": ext.upper() if ext else "-",
                        "الحجم KB": round(len(f.getbuffer()) / 1024, 1),
                    })
                st.dataframe(preview_rows, use_container_width=True, hide_index=True)

            if st.button(
                "📥 حفظ الشحنة وملفاتها",
                type="primary",
                key="save_intake_batch",
                disabled=not bool(intake_files),
            ):
                intake_id = execute("""
                    INSERT INTO shipment_intakes (
                        factory_id, order_id, intake_name, status, private_notes, created_at
                    )
                    VALUES (?, ?, ?, 'Uploaded', ?, CURRENT_TIMESTAMP)
                """, (
                    intake_factory_id,
                    intake_order_map[intake_order_label],
                    intake_name.strip() or f"Shipment Intake {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                    intake_private_notes.strip(),
                ))

                target_dir = SHIPMENT_INBOX_ROOT / str(intake_id)
                target_dir.mkdir(parents=True, exist_ok=True)

                for uploaded in intake_files:
                    safe_name = Path(uploaded.name).name
                    target = target_dir / safe_name
                    if target.exists():
                        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                        target = target_dir / f"{target.stem}_{stamp}{target.suffix}"

                    target.write_bytes(uploaded.getbuffer())
                    ext = target.suffix.lower().replace(".", "")

                    execute("""
                        INSERT INTO shipment_intake_files (
                            intake_id, file_name, file_path, file_ext, detected_type, uploaded_at
                        )
                        VALUES (?, ?, ?, ?, 'Unclassified', CURRENT_TIMESTAMP)
                    """, (
                        intake_id,
                        target.name,
                        str(target),
                        ext,
                    ))

                st.success("تم حفظ الشحنة وملفاتها الأصلية بنجاح.")
                st.info("المرحلة التالية ستضيف المعالجة الذكية للملفات لاستخراج B/L وPI وعدد الحاويات وETD وETA والمبالغ تلقائياً.")
                st.rerun()

            st.divider()
            st.markdown("### 📦 الشحنات المرفوعة من الملفات")

            intake_rows = fetchall("""
                SELECT i.*, f.factory_name, o.pi_no
                FROM shipment_intakes i
                LEFT JOIN factories f ON f.id = i.factory_id
                LEFT JOIN orders o ON o.id = i.order_id
                ORDER BY i.id DESC
            """)

            if not intake_rows:
                st.caption("لا توجد شحنات مرفوعة بهذه الطريقة بعد.")
            else:
                for intake in intake_rows:
                    with st.expander(
                        f"#{intake['id']} — {intake['factory_name'] or '-'} — {intake['intake_name'] or 'Shipment'}"
                    ):
                        st.write(f"**PI:** {intake['pi_no'] or '-'}")
                        st.write(f"**الحالة:** {intake['status'] or 'Uploaded'}")
                        if intake["private_notes"]:
                            st.info(intake["private_notes"])

                        files = fetchall("""
                            SELECT file_name, file_ext, detected_type, uploaded_at
                            FROM shipment_intake_files
                            WHERE intake_id=?
                            ORDER BY id
                        """, (intake["id"],))

                        st.dataframe(
                            [
                                {
                                    "File": x["file_name"],
                                    "Type": (x["file_ext"] or "").upper(),
                                    "Detected As": x["detected_type"] or "Unclassified",
                                    "Uploaded": x["uploaded_at"],
                                }
                                for x in files
                            ],
                            use_container_width=True,
                            hide_index=True,
                        )


                        if st.button("⚙️ تحليل الملفات واستخراج بيانات الشحنة", key=f"analyze_intake_{intake['id']}", type="primary"):
                            analyze_intake_files(intake["id"])
                            st.success("تم تحليل الملفات. راجع البيانات المستخرجة أدناه.")
                            st.rerun()

                        analyzed = fetchone("SELECT * FROM shipment_intakes WHERE id=?", (intake["id"],))
                        if analyzed and analyzed["status"] in ("Analyzed", "Converted"):
                            st.markdown("### 🔎 البيانات المستخرجة")
                            with st.form(f"review_intake_{intake['id']}"):
                                r1, r2, r3 = st.columns(3)
                                ex_pi = r1.text_input("PI No.", value=analyzed["extracted_pi_no"] or "")
                                ex_bl = r2.text_input("B/L No.", value=analyzed["extracted_bl_no"] or "")
                                ex_shipco = r3.text_input("Shipping Company", value=analyzed["extracted_shipping_company"] or "")

                                r4, r5, r6 = st.columns(3)
                                ex_dest = r4.text_input("Destination", value=analyzed["extracted_destination"] or "")
                                ex_count = r5.number_input("عدد الحاويات", min_value=0.0, step=0.5, value=float(analyzed["extracted_containers_count"] or 0))
                                ex_type = r6.selectbox("نوع الحاوية", ["40HQ", "40GP", "20GP", "Other"], index=["40HQ", "40GP", "20GP", "Other"].index(analyzed["extracted_container_type"] or "40HQ") if (analyzed["extracted_container_type"] or "40HQ") in ["40HQ", "40GP", "20GP", "Other"] else 0)

                                r7, r8 = st.columns(2)
                                ex_ci = r7.number_input("CI Amount", min_value=0.0, step=100.0, value=float(analyzed["extracted_ci_amount"] or 0), format="%.2f")
                                ex_shipping = r8.number_input("Shipping Cost", min_value=0.0, step=100.0, value=float(analyzed["extracted_shipping_cost"] or 0), format="%.2f")

                                r9, r10 = st.columns(2)
                                ex_etd = r9.text_input("ETD", value=analyzed["extracted_etd"] or "")
                                ex_eta = r10.text_input("ETA", value=analyzed["extracted_eta"] or "")

                                ex_containers = st.text_area("Container Numbers — رقم بكل سطر", value=analyzed["extracted_container_numbers"] or "")
                                ex_status = st.selectbox("حالة الشحنة", ["Preparing", "Ready", "Booked", "Shipped", "In Transit", "Arrived", "Customs", "Delivered"])

                                create_from_files = st.form_submit_button("✅ اعتماد البيانات وإنشاء الشحنة", type="primary")

                                if create_from_files:
                                    total_amount = float(ex_ci) + float(ex_shipping)
                                    shipment_id = execute("""
                                        INSERT INTO shipments (
                                            factory_id, order_id, shipment_no, bl_no,
                                            shipping_company, destination, container_type,
                                            containers_count, ci_amount, shipping_cost,
                                            total_amount, etd, eta, shipment_status,
                                            payment_status, private_notes, public_notes,
                                            created_at, updated_at
                                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'Not Paid', ?, '', CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                                    """, (
                                        intake["factory_id"], intake["order_id"], intake["intake_name"],
                                        ex_bl.strip(), ex_shipco.strip(), ex_dest.strip(), ex_type,
                                        float(ex_count), float(ex_ci), float(ex_shipping), total_amount,
                                        ex_etd.strip(), ex_eta.strip(), ex_status,
                                        intake["private_notes"] or "",
                                    ))

                                    # If no linked order exists, try to match by PI number within the same factory
                                    matched_order = intake["order_id"]
                                    if not matched_order and ex_pi.strip():
                                        m = fetchone("SELECT id FROM orders WHERE factory_id=? AND pi_no=? ORDER BY id DESC LIMIT 1", (intake["factory_id"], ex_pi.strip()))
                                        if m:
                                            matched_order = m["id"]
                                            execute("UPDATE shipments SET order_id=? WHERE id=?", (matched_order, shipment_id))

                                    # Copy original files into the real shipment documents and keep detected classification
                                    file_rows = fetchall("SELECT * FROM shipment_intake_files WHERE intake_id=? ORDER BY id", (intake["id"],))
                                    shipment_dir = SHIPMENT_UPLOAD_ROOT / str(shipment_id)
                                    shipment_dir.mkdir(parents=True, exist_ok=True)
                                    for fr in file_rows:
                                        source = Path(fr["file_path"])
                                        target = shipment_dir / source.name
                                        if source.exists():
                                            shutil.copy2(source, target)
                                            execute("""
                                                INSERT INTO shipment_documents (shipment_id, document_type, file_name, file_path, uploaded_at)
                                                VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
                                            """, (shipment_id, fr["detected_type"] or "Unclassified", target.name, str(target)))

                                    for cn in [x.strip() for x in ex_containers.splitlines() if x.strip()]:
                                        execute("""
                                            INSERT INTO shipment_containers (shipment_id, container_no, seal_no, notes, created_at)
                                            VALUES (?, ?, '', '', CURRENT_TIMESTAMP)
                                        """, (shipment_id, cn))

                                    execute("""
                                        UPDATE shipment_intakes
                                        SET status='Converted', extracted_pi_no=?, extracted_bl_no=?,
                                            extracted_shipping_company=?, extracted_destination=?,
                                            extracted_container_type=?, extracted_containers_count=?,
                                            extracted_ci_amount=?, extracted_shipping_cost=?,
                                            extracted_etd=?, extracted_eta=?, extracted_container_numbers=?
                                        WHERE id=?
                                    """, (
                                        ex_pi.strip(), ex_bl.strip(), ex_shipco.strip(), ex_dest.strip(), ex_type,
                                        float(ex_count), float(ex_ci), float(ex_shipping), ex_etd.strip(), ex_eta.strip(),
                                        ex_containers.strip(), intake["id"]
                                    ))
                                    st.success(f"تم إنشاء الشحنة #{shipment_id} ونقل ملفاتها تلقائياً إلى مستندات الشحنة.")
                                    st.rerun()

                            if analyzed["analysis_notes"]:
                                st.warning(analyzed["analysis_notes"])

        # ---------------- MANUAL SHIPMENT ----------------
        with tab2:
            st.subheader("إنشاء شحنة جديدة")

            factory_label = st.selectbox(
                "المعمل *",
                list(factory_map.keys()),
                key="shipment_factory_select",
            )
            selected_factory_id = factory_map[factory_label]

            order_rows = fetchall("""
                SELECT id, pi_no, order_status
                FROM orders
                WHERE factory_id=?
                ORDER BY id DESC
            """, (selected_factory_id,))

            order_map = {"بدون ربط بطلبية": None}
            for r in order_rows:
                order_map[f"{r['pi_no'] or 'Order'} — #{r['id']}"] = r["id"]

            with st.form("new_shipment_form", clear_on_submit=True):
                c1, c2, c3 = st.columns(3)
                linked_order_label = c1.selectbox("الطلبية المرتبطة", list(order_map.keys()))
                shipment_no = c2.text_input("رقم الشحنة / Shipment No.")
                bl_no = c3.text_input("B/L No.")

                c4, c5, c6 = st.columns(3)
                shipping_company = c4.text_input("Shipping Company", placeholder="MSC / HPL / CMA / OOCL")
                destination = c5.text_input("Destination", placeholder="UMM QASR / AQABA / MERSIN")
                container_type = c6.selectbox("نوع الحاوية", ["40HQ", "40GP", "20GP", "Other"])

                c7, c8, c9 = st.columns(3)
                containers_count = c7.number_input(
                    "عدد الحاويات",
                    min_value=0.0,
                    step=0.5,
                    format="%.2f"
                )
                ci_amount = c8.number_input(
                    "CI Amount / USD",
                    min_value=0.0,
                    step=100.0,
                    format="%.2f"
                )
                shipping_cost = c9.number_input(
                    "Shipping Cost",
                    min_value=0.0,
                    step=100.0,
                    format="%.2f"
                )

                c10, c11, c12 = st.columns(3)
                etd = c10.date_input("ETD", value=date.today())
                eta = c11.date_input("ETA", value=date.today())
                shipment_status = c12.selectbox(
                    "حالة الشحنة",
                    [
                        "Preparing",
                        "Ready",
                        "Booked",
                        "Shipped",
                        "In Transit",
                        "Arrived",
                        "Customs",
                        "Delivered",
                        "Cancelled",
                    ]
                )

                payment_status = st.selectbox(
                    "حالة الدفع",
                    [
                        "Not Paid",
                        "Deposit Paid",
                        "Partially Paid",
                        "Paid",
                        "Supplier Confirmed",
                    ]
                )

                n1, n2 = st.columns(2)
                private_notes = n1.text_area("🔒 ملاحظات خاصة")
                public_notes = n2.text_area("📝 ملاحظات عامة")

                create_shipment = st.form_submit_button("💾 حفظ الشحنة", type="primary")

                if create_shipment:
                    total_amount = float(ci_amount) + float(shipping_cost)
                    shipment_id = execute("""
                        INSERT INTO shipments (
                            factory_id, order_id, shipment_no, bl_no,
                            shipping_company, destination, container_type,
                            containers_count, ci_amount, shipping_cost,
                            total_amount, etd, eta, shipment_status,
                            payment_status, private_notes, public_notes,
                            created_at, updated_at
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """, (
                        selected_factory_id,
                        order_map[linked_order_label],
                        shipment_no.strip(),
                        bl_no.strip(),
                        shipping_company.strip(),
                        destination.strip(),
                        container_type,
                        float(containers_count),
                        float(ci_amount),
                        float(shipping_cost),
                        total_amount,
                        str(etd),
                        str(eta),
                        shipment_status,
                        payment_status,
                        private_notes.strip(),
                        public_notes.strip(),
                    ))

                    st.success(f"تم إنشاء الشحنة رقم {shipment_id} داخل MIDO.")

        # ---------------- MAIN SHIPMENT LIST ----------------
        with tab3:
            shipments = fetchall("""
                SELECT
                    s.*,
                    f.factory_name,
                    f.brand_name,
                    o.pi_no
                FROM shipments s
                LEFT JOIN factories f ON f.id = s.factory_id
                LEFT JOIN orders o ON o.id = s.order_id
                ORDER BY s.id DESC
            """)

            if not shipments:
                st.info("لا توجد مشحونات بعد.")
            else:
                st.subheader("المشحونات الرئيسية")

                f1, f2, f3 = st.columns(3)
                search = f1.text_input(
                    "🔎 بحث",
                    placeholder="PI / B-L / المعمل / شركة الشحن",
                    key="shipment_search"
                )
                status_filter = f2.selectbox(
                    "فلترة حسب الحالة",
                    ["الكل", "Preparing", "Ready", "Booked", "Shipped", "In Transit", "Arrived", "Customs", "Delivered", "Cancelled"]
                )
                payment_filter = f3.selectbox(
                    "فلترة حسب الدفع",
                    ["الكل", "Not Paid", "Deposit Paid", "Partially Paid", "Paid", "Supplier Confirmed"]
                )

                filtered = []
                for s in shipments:
                    haystack = " ".join([
                        str(s["shipment_no"] or ""),
                        str(s["bl_no"] or ""),
                        str(s["pi_no"] or ""),
                        str(s["factory_name"] or ""),
                        str(s["brand_name"] or ""),
                        str(s["shipping_company"] or ""),
                        str(s["destination"] or ""),
                    ]).lower()

                    ok_search = (not search.strip()) or (search.lower().strip() in haystack)
                    ok_status = status_filter == "الكل" or s["shipment_status"] == status_filter
                    ok_payment = payment_filter == "الكل" or s["payment_status"] == payment_filter

                    if ok_search and ok_status and ok_payment:
                        filtered.append(s)

                total_conts = sum(float(s["containers_count"] or 0) for s in filtered)
                total_ci = sum(float(s["ci_amount"] or 0) for s in filtered)
                total_shipping = sum(float(s["shipping_cost"] or 0) for s in filtered)

                m1, m2, m3, m4 = st.columns(4)
                m1.metric("عدد المشحونات", len(filtered))
                m2.metric("إجمالي الحاويات", f"{total_conts:g}")
                m3.metric("CI Amount", f"${total_ci:,.2f}")
                m4.metric("Shipping Cost", f"${total_shipping:,.2f}")

                table_rows = []
                for s in filtered:
                    table_rows.append({
                        "PI No.": s["pi_no"] or "-",
                        "Factory": s["factory_name"] or "-",
                        "Shipment No.": s["shipment_no"] or "-",
                        "BL": s["bl_no"] or "-",
                        "Ship Com.": s["shipping_company"] or "-",
                        "Destination": s["destination"] or "-",
                        "CONTS": s["containers_count"] or 0,
                        "Type": s["container_type"] or "40HQ",
                        "CI Amount/USD": s["ci_amount"] or 0,
                        "Shipping": s["shipping_cost"] or 0,
                        "Total Amount": s["total_amount"] or 0,
                        "ETD": s["etd"] or "-",
                        "ETA": s["eta"] or "-",
                        "State": s["shipment_status"] or "-",
                        "Payment": s["payment_status"] or "-",
                    })

                st.dataframe(table_rows, use_container_width=True, hide_index=True)

                st.divider()
                st.subheader("تفاصيل كل شحنة")

                for s in filtered:
                    title = (
                        f"#{s['id']} — {s['factory_name'] or '-'} — "
                        f"{s['pi_no'] or 'No PI'} — {s['bl_no'] or 'No BL'}"
                    )

                    with st.expander(title):
                        a1, a2, a3, a4 = st.columns(4)
                        a1.metric("الحاويات", f"{float(s['containers_count'] or 0):g}")
                        a2.metric("CI Amount", f"${float(s['ci_amount'] or 0):,.2f}")
                        a3.metric("Shipping", f"${float(s['shipping_cost'] or 0):,.2f}")
                        a4.metric("Total", f"${float(s['total_amount'] or 0):,.2f}")

                        d1, d2, d3 = st.columns(3)
                        d1.write(f"**Factory:** {s['factory_name'] or '-'}")
                        d1.write(f"**PI No.:** {s['pi_no'] or '-'}")
                        d1.write(f"**Shipment No.:** {s['shipment_no'] or '-'}")

                        d2.write(f"**B/L:** {s['bl_no'] or '-'}")
                        d2.write(f"**Shipping Company:** {s['shipping_company'] or '-'}")
                        d2.write(f"**Destination:** {s['destination'] or '-'}")

                        d3.write(f"**ETD:** {s['etd'] or '-'}")
                        d3.write(f"**ETA:** {s['eta'] or '-'}")
                        d3.write(f"**Status:** {s['shipment_status'] or '-'}")

                        n1, n2 = st.columns(2)
                        n1.info(s["private_notes"] or "لا توجد ملاحظات خاصة.")
                        n2.info(s["public_notes"] or "لا توجد ملاحظات عامة.")

                        st.markdown("### 🚛 أرقام الحاويات")
                        with st.form(f"container_form_{s['id']}", clear_on_submit=True):
                            cc1, cc2, cc3 = st.columns(3)
                            container_no = cc1.text_input("Container No.")
                            seal_no = cc2.text_input("Seal No.")
                            container_note = cc3.text_input("ملاحظة")

                            if st.form_submit_button("➕ إضافة الحاوية"):
                                if not container_no.strip():
                                    st.error("اكتب رقم الحاوية.")
                                else:
                                    execute("""
                                        INSERT INTO shipment_containers (
                                            shipment_id, container_no, seal_no, notes, created_at
                                        )
                                        VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
                                    """, (
                                        s["id"],
                                        container_no.strip(),
                                        seal_no.strip(),
                                        container_note.strip(),
                                    ))
                                    st.success("تمت إضافة الحاوية.")
                                    st.rerun()

                        cont_rows = fetchall("""
                            SELECT container_no, seal_no, notes
                            FROM shipment_containers
                            WHERE shipment_id=?
                            ORDER BY id
                        """, (s["id"],))

                        if cont_rows:
                            st.dataframe(
                                [
                                    {
                                        "Container No.": r["container_no"],
                                        "Seal No.": r["seal_no"],
                                        "Notes": r["notes"],
                                    }
                                    for r in cont_rows
                                ],
                                use_container_width=True,
                                hide_index=True,
                            )
                        else:
                            st.caption("لا توجد أرقام حاويات مضافة بعد.")

                        st.markdown("### 📁 مستندات الشحنة")
                        shipment_doc_type = st.selectbox(
                            "نوع المستند",
                            [
                                "Commercial Invoice",
                                "Packing List",
                                "Certificate of Origin (CO)",
                                "COC",
                                "Bill of Lading",
                                "QR Code",
                                "Other",
                            ],
                            key=f"shipment_doc_type_{s['id']}",
                        )

                        shipment_files = st.file_uploader(
                            "ارفع ملف أو أكثر",
                            accept_multiple_files=True,
                            type=["pdf", "xlsx", "xls", "csv", "docx", "doc", "png", "jpg", "jpeg", "webp"],
                            key=f"shipment_upload_{s['id']}",
                        )

                        if shipment_files and st.button(
                            "⬆️ حفظ مستندات الشحنة",
                            key=f"save_shipment_docs_{s['id']}",
                            type="primary",
                        ):
                            saved = save_uploaded_files(
                                shipment_files,
                                SHIPMENT_UPLOAD_ROOT / str(s["id"]),
                                "shipment_documents",
                                "shipment_id",
                                s["id"],
                                shipment_doc_type,
                            )
                            st.success(f"تم حفظ {saved} ملف/ملفات.")
                            st.rerun()

                        docs = fetchall("""
                            SELECT document_type, file_name, uploaded_at
                            FROM shipment_documents
                            WHERE shipment_id=?
                            ORDER BY id DESC
                        """, (s["id"],))

                        if docs:
                            st.dataframe(
                                [
                                    {
                                        "Document": d["document_type"],
                                        "File": d["file_name"],
                                        "Uploaded": d["uploaded_at"],
                                    }
                                    for d in docs
                                ],
                                use_container_width=True,
                                hide_index=True,
                            )
                        else:
                            st.caption("لا توجد مستندات لهذه الشحنة بعد.")

                        st.markdown("### ✏️ تحديث حالة الشحنة")
                        with st.form(f"update_shipment_{s['id']}"):
                            u1, u2, u3 = st.columns(3)
                            new_status = u1.selectbox(
                                "الحالة",
                                ["Preparing", "Ready", "Booked", "Shipped", "In Transit", "Arrived", "Customs", "Delivered", "Cancelled"],
                                index=["Preparing", "Ready", "Booked", "Shipped", "In Transit", "Arrived", "Customs", "Delivered", "Cancelled"].index(s["shipment_status"]) if s["shipment_status"] in ["Preparing", "Ready", "Booked", "Shipped", "In Transit", "Arrived", "Customs", "Delivered", "Cancelled"] else 0,
                                key=f"status_{s['id']}"
                            )
                            new_payment = u2.selectbox(
                                "الدفع",
                                ["Not Paid", "Deposit Paid", "Partially Paid", "Paid", "Supplier Confirmed"],
                                index=["Not Paid", "Deposit Paid", "Partially Paid", "Paid", "Supplier Confirmed"].index(s["payment_status"]) if s["payment_status"] in ["Not Paid", "Deposit Paid", "Partially Paid", "Paid", "Supplier Confirmed"] else 0,
                                key=f"payment_{s['id']}"
                            )
                            new_containers = u3.number_input(
                                "عدد الحاويات",
                                min_value=0.0,
                                step=0.5,
                                value=float(s["containers_count"] or 0),
                                key=f"containers_{s['id']}"
                            )

                            new_private = st.text_area(
                                "🔒 ملاحظات خاصة",
                                value=s["private_notes"] or "",
                                key=f"private_{s['id']}"
                            )
                            new_public = st.text_area(
                                "📝 ملاحظات عامة",
                                value=s["public_notes"] or "",
                                key=f"public_{s['id']}"
                            )

                            if st.form_submit_button("💾 حفظ التحديث"):
                                execute("""
                                    UPDATE shipments
                                    SET shipment_status=?,
                                        payment_status=?,
                                        containers_count=?,
                                        private_notes=?,
                                        public_notes=?,
                                        updated_at=CURRENT_TIMESTAMP
                                    WHERE id=?
                                """, (
                                    new_status,
                                    new_payment,
                                    float(new_containers),
                                    new_private.strip(),
                                    new_public.strip(),
                                    s["id"],
                                ))
                                st.success("تم تحديث الشحنة.")
                                st.rerun()

elif page == "الدفعات":
    st.title("💰 الدفعات")
    st.info("سيتم ربط الدفعات بالطلبيات والمشحونات.")

elif page == "المستندات":
    st.title("📁 المستندات")
    st.info("Invoice / Packing List / CO / COC / B/L / QR Code")

elif page == "نظام Access القديم":
    st.title("🗃️ نظام Access القديم")
    st.caption("هذه الصفحة تربط MIDO بهيكل قاعدة برنامج Access القديم حتى ننقل أقسامه ونطوّرها داخل MIDO.")

    if not access_db_available():
        st.error(
            "ملف mido_access_structure_only.db غير موجود بجانب app.py. "
            "ارفع قاعدة Access الفارغة إلى نفس مجلد GitHub."
        )
    else:
        tables = access_table_names()

        a1, a2, a3 = st.columns(3)
        a1.metric("جداول Access", len(tables))
        a2.metric("قاعدة Access", "مرتبطة ✅")
        a3.metric("الحالة", "Structure Only")

        st.info(
            "القاعدة الموجودة حالياً هي نسخة الهيكل فقط بدون البيانات القديمة. "
            "هذا مقصود حتى تبقى خفيفة على GitHub. سنبني واجهات MIDO فوق هذه الجداول، "
            "ثم ننقل البيانات التي تحتاجها بطريقة منفصلة وآمنة."
        )

        categories = {
            "👥 العملاء والحسابات": [
                "جدول_الزبائن", "حساب_عميل", "حساب_سابق", "حساب_سابق_دائن",
                "تسديدات_العملاء", "رصيد_افتتاحي_عملاء"
            ],
            "📦 المواد والمخزون": [
                "جدول_المواد", "اصناف", "محزن", "store", "stores",
                "كميات_الارصدة_الافتتاحية", "الرصيد_الاقتتاحي"
            ],
            "🧾 المبيعات": [
                "جدول_البيع", "جدول_التفاصيل", "ارباح", "ord", "order"
            ],
            "🛒 المشتريات والمجهزين": [
                "مشتريات", "حساب_مجهز_كلي", "تسديدات_مشتريات_مواد"
            ],
            "💰 الصندوق والصيرفة": [
                "صندوق", "حساب_صيرفة", "مكاتب_صيرفة", "تسديدات_الصرفيات", "راسمال"
            ],
            "🏢 المكاتب والأطراف": [
                "مكاتب", "حساب_مكتب", "متعهدين", "مندوبين", "ماركة", "ماركة1"
            ],
        }

        st.subheader("الأقسام الموجودة في برنامج Access")
        for title, names in categories.items():
            existing = [name for name in names if name in tables]
            with st.expander(f"{title} — {len(existing)} جدول"):
                if existing:
                    for name in existing:
                        cols = access_table_columns(name)
                        st.write(f"**{name}** — {len(cols)} حقل")
                        if cols:
                            st.caption(" • ".join(cols[:12]) + (" ..." if len(cols) > 12 else ""))
                else:
                    st.caption("لا توجد جداول من هذه المجموعة في النسخة الحالية.")

        st.divider()
        st.subheader("استعراض جميع جداول Access")
        if tables:
            selected_table = st.selectbox("اختر جدول", tables, key="access_table_selector")
            cols = access_table_columns(selected_table)
            row_count = access_table_count(selected_table)

            b1, b2 = st.columns(2)
            b1.metric("عدد الحقول", len(cols))
            b2.metric("عدد السجلات الحالية", row_count)

            if cols:
                st.write("**الحقول:**")
                st.code("\n".join(cols), language=None)

            if row_count:
                rows = access_table_rows(selected_table, 200)
                st.dataframe([dict(r) for r in rows], use_container_width=True, hide_index=True)
            else:
                st.warning("هذا الجدول فارغ حالياً لأننا رفعنا نسخة الهيكل فقط.")
        else:
            st.warning("لم أجد جداول Access قابلة للعرض.")

elif page == "التقارير":
    st.title("📄 التقارير")
    st.info("سيتم بناء تقارير PDF وExcel للمشحونات.")
